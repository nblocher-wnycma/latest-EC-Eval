import json
import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SECTIONS = [
    "Basic Info",
    "Barnyard",
    "Clean Stormwater",
    "Livestock Manure Storage",
    "Pasture",
    "Animal Mortalities",
    "Fuel / Petroleum",
    "Air Quality",
    "Driveways",
    "Feed Storage",
]

SECTION_SYSTEM_NAMES = {
    "Barnyard": "Livestock Heavy Use Area Runoff Management System",
    "Clean Stormwater": "Short-Term Waste Collection and Transfer System",
    "Livestock Manure Storage": "Manure and Agricultural Waste Treatment Systems",
    "Pasture": "Livestock Heavy Use Area Runoff Management System",
    "Animal Mortalities": "Composting System – Animal",
    "Fuel / Petroleum": "Petroleum and Oil Products Storage System",
    "Air Quality": "Air Quality",
    "Driveways": "Access Control System",
    "Feed Storage": "Silage Leachate Control and Treatment System",
}

QUICK_BUILD_SECTIONS = [s for s in SECTIONS if s != "Basic Info"]


BMP_SYSTEM_COMPONENTS = {
    "Access Control System": [
        ("Access Control", "472"), ("Fence", "382"), ("Trails and Walkways", "575"),
        ("Access Road", "560"), ("Watering Facility", "614"), ("Heavy Use Area Protection", "561"),
        ("Livestock Pipeline", "516"), ("Pond", "378"), ("Pumping Plant", "533"),
        ("Spring Development", "574"), ("Structure for Water Control", "587"), ("Stream Crossing", "578"),
        ("Water Well", "642"), ("Field Border", "386"),
    ],
    "Livestock Heavy Use Area Runoff Management System": [
        ("Heavy Use Area Protection", "561"), ("Access Control", "472"), ("Trails and Walkways", "575"),
        ("Fence", "382"), ("Access Road", "560"), ("Roof Runoff Structure", "558"),
        ("Roofs and Covers", "367"), ("Conservation Cover", "327"), ("Critical Area Planting", "342"),
        ("Vegetated Treatment Area", "635"), ("Diversion", "362"), ("Grassed Waterway", "412"),
        ("Lined Waterway or Outlet", "468"), ("Pumping Plant", "533"), ("Sediment Basin", "350"),
        ("Subsurface Drain", "606"), ("Underground Outlet", "620"), ("Waste Separation Facility", "632"),
        ("Waste Storage Facility", "313"), ("Waste Transfer", "634"),
        ("Water and Sediment Control Basin (WASCOB)", "638"), ("Watering Facility", "614"),
    ],
    "Short-Term Waste Collection and Transfer System": [
        ("Waste Storage Facility", "313"),
        ("NYS SWCC Short-Term Waste Collection and Transfer System – Modified NRCS CPS Waste Storage Facility", "313"),
        ("Waste Transfer", "634"), ("Waste Separation Facility", "632"), ("Waste Treatment", "629"),
        ("Waste Facility Closure", "360"), ("Nutrient Management", "590"), ("Pumping Plant", "533"),
    ],
    "Manure and Agricultural Waste Treatment Systems": [
        ("Waste Separation Facility", "632"), ("Waste Storage Facility", "313"), ("Waste Transfer", "634"),
        ("Waste Treatment", "629"), ("Anaerobic Digester", "366"), ("Composting Facility", "317"),
        ("Pumping Plant", "533"), ("Roofs and Covers", "367"), ("Heavy Use Area Protection", "561"),
    ],
    "Composting System – Animal": [
        ("Animal Mortality Facility", "316"), ("Composting Facility", "317"), ("Access Road", "560"),
        ("Diversion", "362"), ("Heavy Use Area Protection", "561"), ("Roofs and Covers", "367"),
        ("Vegetated Treatment Area", "635"), ("Waste Storage Facility", "313"), ("Waste Transfer", "634"),
    ],
    "Petroleum and Oil Products Storage System": [
        ("On-Farm Secondary Containment Facility", "319"), ("Access Road", "560"),
        ("Access Control", "472"), ("Heavy Use Area Protection", "561"), ("Conservation Cover", "327"),
    ],
    "Silage Leachate Control and Treatment System": [
        # This system was requested for the Feed Storage tab. Edit the standards override if your local SWCC list uses different components.
        ("Vegetated Treatment Area", "635"), ("Waste Storage Facility", "313"), ("Waste Transfer", "634"),
        ("Roof Runoff Structure", "558"), ("Underground Outlet", "620"), ("Subsurface Drain", "606"),
        ("Heavy Use Area Protection", "561"), ("Critical Area Planting", "342"),
    ],
}

# All common NRCS-NY component practices from the BMP systems catalogue used for the Quick Build search.
# Keep this list broad so each section can search and add multiple planned practices.
ALL_BMP_STANDARDS = [
    ("Access Control", "472"),
    ("Access Road", "560"),
    ("Agrichemical Handling Facility", "309"),
    ("Anaerobic Digester", "366"),
    ("Animal Mortality Facility", "316"),
    ("Brush Management", "314"),
    ("Composting Facility", "317"),
    ("Conservation Cover", "327"),
    ("Conservation Crop Rotation", "328"),
    ("Contour Orchard and Other Perennial Crops", "331"),
    ("Critical Area Planting", "342"),
    ("Diversion", "362"),
    ("Fence", "382"),
    ("Field Border", "386"),
    ("Filter Strip", "393"),
    ("Forest Farming", "379"),
    ("Forest Stand Improvement", "666"),
    ("Forest Trails and Landings", "655"),
    ("Forage Harvest Management", "511"),
    ("Grazing Land Mechanical Treatment", "548"),
    ("Grassed Waterway", "412"),
    ("Ground Water Testing", "355"),
    ("Heavy Use Area Protection", "561"),
    ("Hedgerow Planting", "422"),
    ("Herbaceous Weed Treatment", "315"),
    ("Irrigation Pipeline", "430"),
    ("Irrigation Reservoir", "436"),
    ("Irrigation System, Microirrigation", "441"),
    ("Irrigation System, Surface and Subsurface", "443"),
    ("Irrigation Water Management", "449"),
    ("Lined Waterway or Outlet", "468"),
    ("Livestock Pipeline", "516"),
    ("Nutrient Management", "590"),
    ("NYS SWCC Short-Term Waste Collection and Transfer System – Modified NRCS CPS Waste Storage Facility", "313"),
    ("Obstruction Removal", "500"),
    ("On-Farm Secondary Containment Facility", "319"),
    ("Pasture and Hay Planting", "512"),
    ("Pest Management Conservation System", "595"),
    ("Pond", "378"),
    ("Prescribed Grazing", "528"),
    ("Pumping Plant", "533"),
    ("Riparian Forest Buffer", "391"),
    ("Riparian Herbaceous Cover", "390"),
    ("Roof Runoff Structure", "558"),
    ("Roofs and Covers", "367"),
    ("Sediment Basin", "350"),
    ("Spring Development", "574"),
    ("Sprinkler System", "442"),
    ("Stream Crossing", "578"),
    ("Structure for Water Control", "587"),
    ("Subsurface Drain", "606"),
    ("Terrace", "600"),
    ("Trails and Walkways", "575"),
    ("Tree/Shrub Establishment", "612"),
    ("Tree/Shrub Site Preparation", "490"),
    ("Underground Outlet", "620"),
    ("Vegetated Treatment Area", "635"),
    ("Waste Facility Closure", "360"),
    ("Waste Separation Facility", "632"),
    ("Waste Storage Facility", "313"),
    ("Waste Transfer", "634"),
    ("Waste Treatment", "629"),
    ("Water and Sediment Control Basin (WASCOB)", "638"),
    ("Water Well", "642"),
    ("Watering Facility", "614"),
    ("Windbreak-Shelterbelt Establishment and Renovation", "380"),
]

ALL_BMP_CHOICES = [f"{practice} ({code})" for practice, code in ALL_BMP_STANDARDS]
ALL_BMP_CHOICES = sorted(dict.fromkeys(ALL_BMP_CHOICES), key=lambda x: x.lower())

def all_bmp_dropdown_choices():
    return ["No structural BMP - continue O&M"] + ALL_BMP_CHOICES

def _practice_name_from_line(line):
    """Return just the practice name/code from either a plain practice line or a standards-table line."""
    line = (line or "").strip()
    if "|" in line:
        return line.split("|", 1)[0].strip()
    return line

def practice_lines_to_sentence(text):
    lines = [_practice_name_from_line(ln) for ln in (text or "").splitlines() if ln.strip()]
    lines = [ln for ln in lines if ln and ln != "No structural BMP - continue O&M"]
    if not lines:
        return ""
    if len(lines) == 1:
        return lines[0]
    if len(lines) == 2:
        return lines[0] + " and " + lines[1]
    return ", ".join(lines[:-1]) + ", and " + lines[-1]

def standards_for_bmp_choices(choices_text, system_hint=None):
    lines = [ln.strip() for ln in (choices_text or "").replace(";", "\n").splitlines() if ln.strip()]
    output = []
    seen = set()
    for line in lines:
        if "|" in line:
            out_line = line.strip()
            if out_line and out_line not in seen:
                output.append(out_line)
                seen.add(out_line)
            continue
        block = standards_for_bmp_choice(line, system_hint)
        for out_line in block.splitlines():
            out_line = out_line.strip()
            if out_line and out_line not in seen:
                output.append(out_line)
                seen.add(out_line)
    return "\n".join(output) + ("\n" if output else "")

def quick_bmp_list_for(section):
    """Return Quick Build planned-action choices with the requested BMP system and component practices."""
    system = SECTION_SYSTEM_NAMES.get(section, section)
    choices = []
    if system and system != "Air Quality":
        choices.append(system)
        for practice, code in BMP_SYSTEM_COMPONENTS.get(system, []):
            choices.append(f"{practice} ({code})")
    defaults = {
        "Air Quality": ["Continue O&M and proper manure application", "Dust/odor management", "Manure management changes"],
    }
    choices.extend(defaults.get(section, []))
    choices.append("No structural BMP - continue O&M")
    # Remove duplicates while preserving order.
    seen = set()
    out = []
    for c in choices:
        if c not in seen:
            out.append(c)
            seen.add(c)
    return out

def standards_for_bmp_choice(choice, system_hint=None):
    """Create editable standard lines for a selected BMP system or component practice."""
    choice = (choice or "").strip()
    if not choice or choice == "No structural BMP - continue O&M":
        return ""
    if choice in BMP_SYSTEM_COMPONENTS:
        return "".join([f"{practice} ({code})|[item #]|{practice}|[units]\n" for practice, code in BMP_SYSTEM_COMPONENTS[choice]])
    # Component choice, e.g. Heavy Use Area Protection (561)
    if choice.endswith(")") and "(" in choice:
        desc = choice.rsplit("(", 1)[0].strip()
        code = choice.rsplit("(", 1)[1].rstrip(")").strip()
        if code.isdigit():
            return f"{desc} ({code})|[item #]|{desc}|[units]\n"
    if system_hint and system_hint in BMP_SYSTEM_COMPONENTS:
        for practice, code in BMP_SYSTEM_COMPONENTS[system_hint]:
            if practice.lower() in choice.lower() or f"({code})" in choice:
                return f"{practice} ({code})|[item #]|{practice}|[units]\n"
    return ""

QUICK_BUILD_OPTIONS = {
    "Barnyard": {
        "conditions": [
            "Muddy/denuded barnyard or sacrifice lot",
            "Heavy traffic feeding or loafing area",
            "Manure accumulation and contaminated runoff",
            "Ditch or watercourse receiving barnyard runoff",
            "No current resource concern"
        ],
        "bmps": [
            "Covered heavy use area",
            "Concrete heavy use area",
            "Gravel heavy use area",
            "Restricted access and reseeding",
            "No structural BMP - continue O&M"
        ],
    },
    "Clean Stormwater": {
        "conditions": [
            "Roof runoff enters livestock use area",
            "Clean water contacts manure or temporary piles",
            "Downspout/outlet discharges through degraded area",
            "Upslope clean water enters production area",
            "No current resource concern"
        ],
        "bmps": [
            "Roof runoff structure and underground outlet",
            "Gutters and downspouts",
            "Diversion",
            "Outlet stabilization",
            "No structural BMP - continue O&M"
        ],
    },
    "Livestock Manure Storage": {
        "conditions": [
            "Temporary manure piles with leachate/runoff concern",
            "No permanent manure storage",
            "Storage undersized or poorly located",
            "Clean water contacts manure storage area",
            "No current resource concern"
        ],
        "bmps": [
            "Covered dry stack manure storage",
            "Waste storage facility",
            "Stacking pad",
            "Roof/gutter/outlet improvements",
            "No structural BMP - continue O&M"
        ],
    },
    "Pasture": {
        "conditions": [
            "Continuous grazing with overgrazed/denuded areas",
            "Livestock concentration around gates/water/feed",
            "Poor livestock distribution",
            "Stream or sensitive area access",
            "No current resource concern"
        ],
        "bmps": [
            "Prescribed grazing system",
            "Fence and paddock divisions",
            "Watering facility and livestock pipeline",
            "Trails and walkways / laneway",
            "No structural BMP - continue O&M"
        ],
    },
    "Animal Mortalities": {
        "conditions": ["Burial follows NYS requirements", "Composting follows NYS requirements", "Mortality handling concern", "No current resource concern"],
        "bmps": ["Continue current method and O&M", "Improve mortality site/handling", "Composting facility", "No structural BMP - continue O&M"],
    },
    "Fuel / Petroleum": {
        "conditions": ["No petroleum storage on farm", "Petroleum storage present with no observed concern", "Petroleum storage/spill containment concern", "No current resource concern"],
        "bmps": ["Continue O&M", "Follow DEC tank guidance", "Add containment/spill protection", "No structural BMP - continue O&M"],
    },
    "Air Quality": {
        "conditions": ["Solid manure/no liquid storage with minimal odor", "Dust or odor concern", "Liquid manure odor/ammonia concern", "No current resource concern"],
        "bmps": ["Continue O&M and proper manure application", "Dust/odor management", "Manure management changes", "No structural BMP - continue O&M"],
    },
    "Driveways": {
        "conditions": ["Driveways maintained with minor manure tracking", "Rutted/muddy driveway", "Runoff carries manure/sediment from driveway", "No current resource concern"],
        "bmps": ["Add gravel and maintain driveway", "Clean manure from driveway", "Grade/repair driveway", "No structural BMP - continue O&M"],
    },
    "Feed Storage": {
        "conditions": ["Hay only/no silage storage", "Feed stored with no observed concern", "Silage/feed leachate concern", "Spoiled feed/runoff concern", "No current resource concern"],
        "bmps": ["Continue O&M", "Improve feed storage runoff control", "Clean up spoiled feed", "No structural BMP - continue O&M"],
    },
}

# Replace the original quick-build planned-action dropdowns with the requested BMP systems and their components.
for _sec in QUICK_BUILD_OPTIONS:
    QUICK_BUILD_OPTIONS[_sec]["bmps"] = quick_bmp_list_for(_sec)

EXPANDED_EXISTING_CONDITIONS = {
    "Barnyard": [
        "Muddy, denuded, or compacted surface conditions.",
        "Lack of vegetative cover in livestock concentration areas.",
        "Manure accumulation in areas exposed to rainfall.",
        "Clean stormwater contacting manure, bedding, or contaminated lot surfaces.",
        "Contaminated runoff leaving the barnyard or heavy use area.",
        "Runoff carrying manure, nutrients, pathogens, or sediment toward a ditch, wetland, stream, pond, or road ditch.",
        "Livestock traffic causing soil disturbance, rutting, or erosion.",
        "Undersized barnyard or sacrifice lot for the number of animals using it.",
        "Poor drainage or ponding in the livestock use area.",
        "Manure tracking from the barnyard onto driveways, pastures, or nearby fields.",
        "Livestock congregation around gates, feeders, waterers, or barn entrances causing bare soil.",
        "High water table, hydric soils, floodplain, aquifer, or wetland proximity increasing water quality risk.",
        "No current resource concern."
    ],
    "Clean Stormwater": [
        "Roof runoff entering the barnyard, heavy use area, manure pile, feed area, or livestock lot.",
        "Clean water becoming contaminated by contacting manure, bedding, feed waste, or bare lot surfaces.",
        "Upslope runoff entering the production area.",
        "Missing, damaged, or undersized gutters and downspouts.",
        "Downspouts discharging directly into livestock use areas.",
        "Concentrated clean water flow causing erosion.",
        "Outlet pipes discharging without rock protection or stable vegetated outlet.",
        "Stormwater infiltrating temporary manure piles or manure storage areas.",
        "Clean water increasing the volume of contaminated runoff that must be managed.",
        "Poor separation between clean water and agricultural waste areas.",
        "No current resource concern."
    ],
    "Livestock Manure Storage": [
        "No permanent manure storage on the farmstead.",
        "Temporary manure piles exposed to precipitation.",
        "Manure leachate or runoff leaving the pile area.",
        "Clean water contacting manure piles or storage areas.",
        "Storage located too close to wetlands, streams, ditches, wells, floodplains, or other sensitive areas.",
        "Storage area lacking a roof, cover, curb, pad, or runoff control.",
        "Manure stored directly on soil without adequate containment.",
        "Storage capacity is not adequate for winter or adverse weather periods.",
        "Manure must be spread during unsuitable weather or field conditions because of limited storage.",
        "Saturated manure piles or ponding around storage.",
        "Potential groundwater concern due to shallow depth to water table or permeable soils.",
        "Human sanitary waste not clearly separated from the agricultural waste system.",
        "No current resource concern."
    ],
    "Pasture": [
        "Overgrazing resulting in bare soil or weak forage stands.",
        "Continuous grazing without adequate rest periods.",
        "Livestock congregation around gates, feeders, water tanks, shade, or barn access.",
        "Soil compaction in high-use areas.",
        "Erosion or concentrated flow paths through pasture.",
        "Streambank damage or livestock access to streams, ditches, wetlands, or ponds.",
        "Manure deposition too close to watercourses or wells.",
        "Poor livestock distribution across the pasture.",
        "Lack of fencing or paddock divisions.",
        "Inadequate water distribution causing animals to concentrate in one area.",
        "Laneways running up and down slopes, causing gullies or runoff channels.",
        "Pasture runoff carrying sediment, nutrients, or pathogens off-site.",
        "Pasture areas that remain wet, muddy, or unvegetated for long periods.",
        "No current resource concern."
    ],
    "Animal Mortalities": [
        "No designated mortality management area.",
        "Mortalities not handled within the required timeframe.",
        "Burial, composting, or storage occurring too close to wells, streams, ditches, wetlands, floodplains, or property boundaries.",
        "Shallow burial or improper cover.",
        "Scavenger access to carcasses.",
        "Leachate from mortality handling area reaching soil or water resources.",
        "Composting area lacks proper base, cover material, drainage, or clean water diversion.",
        "Rendering pickup area is poorly located or exposed to runoff.",
        "No backup plan for mortality management if rendering is delayed.",
        "Mortality handling does not follow state and local requirements.",
        "No current resource concern."
    ],
    "Fuel / Petroleum": [
        "Fuel tanks lacking secondary containment.",
        "Single-wall tanks or unprotected tanks in poor condition.",
        "No spill kit or spill response plan.",
        "Evidence of leaks, staining, or spills near fuel storage.",
        "Tanks located near wells, ditches, streams, wetlands, drains, or traffic areas.",
        "No overfill protection.",
        "No protection from vehicle or equipment damage.",
        "Improper storage of oil, hydraulic fluid, waste oil, or other petroleum products.",
        "Containers stored outside without cover or containment.",
        "Potential for petroleum runoff to leave the farmstead.",
        "No current resource concern."
    ],
    "Air Quality": [
        "Odor concerns from manure storage, spreading, composting, or livestock concentration areas.",
        "Dust from driveways, animal lots, feed handling, or field operations.",
        "Ammonia concerns from manure accumulation or liquid manure systems.",
        "Methane concerns from liquid manure storage or anaerobic conditions.",
        "Manure applied during windy conditions near neighbors or public areas.",
        "Inadequate manure incorporation where appropriate.",
        "Poor ventilation in animal housing areas.",
        "Excessive bedding, feed waste, or manure buildup contributing to odor.",
        "Smoke, burning, or improper waste disposal.",
        "Potential nuisance concern for nearby residences, roads, or public areas.",
        "No current resource concern."
    ],
    "Driveways": [
        "Rutted, muddy, or unstable driveway surface.",
        "Manure tracked onto driveways from livestock areas or equipment.",
        "Runoff carrying manure, sediment, or nutrients from the driveway.",
        "Driveway erosion or concentrated flow.",
        "Poor grading causing ponding or water movement through contaminated areas.",
        "Lack of gravel or adequate surface stabilization.",
        "Driveway runoff connected to a ditch, road ditch, stream, wetland, or culvert.",
        "Equipment traffic spreading manure or soil outside the production area.",
        "Access road too narrow, steep, or poorly drained for regular farm use.",
        "Sediment leaving the site from vehicle traffic areas.",
        "No current resource concern."
    ],
    "Feed Storage": [
        "Silage leachate or feed runoff leaving the storage area.",
        "Spoiled feed exposed to rainfall.",
        "Feed stored directly on soil in a way that creates runoff or leachate.",
        "Feed storage located near ditches, streams, wetlands, wells, or drainageways.",
        "Clean stormwater contacting feed waste.",
        "Lack of collection or treatment for leachate.",
        "Bunker, bag, baleage, or feed pad runoff not controlled.",
        "Wildlife or livestock spreading spoiled feed into drainage areas.",
        "Muddy conditions around feed storage or feeding locations.",
        "Feed waste contributing nutrients or organic material to runoff.",
        "No current resource concern."
    ],
}

# Replace the shorter Quick Build condition lists with the more complete checklist-style options.
for _sec, _conds in EXPANDED_EXISTING_CONDITIONS.items():
    if _sec in QUICK_BUILD_OPTIONS:
        QUICK_BUILD_OPTIONS[_sec]["conditions"] = _conds

BMP_STANDARD_TEMPLATES = {
    "Covered heavy use area": "Heavy Use Area Protection (561)|1d|Concrete pad with curbs|[units]\nRoof & Covers (367)|1e|Wood roof|[units]\nRoof Runoff Structure (558)|1f|Gutter & downspout|[units]\nUnderground Outlet (620)|1g|Clean water outlet|[units]\n",
    "Concrete heavy use area": "Heavy Use Area Protection (561)|1a|Concrete heavy use area|[units]\n",
    "Gravel heavy use area": "Heavy Use Area Protection (561)|1a|Gravel heavy use area|[units]\n",
    "Restricted access and reseeding": "Access Control (472)|1a|Control livestock access|[units]\nFence (382)|1b|Fence restricted area|[units]\nCritical Area Planting (342)|1c|Seed/revegetate abused area|[units]\n",
    "Roof runoff structure and underground outlet": "Roof Runoff Structure (558)|1a|Gutter & downspout|[units]\nUnderground Outlet (620)|1b|Clean water outlet|[units]\n",
    "Gutters and downspouts": "Roof Runoff Structure (558)|1a|Gutter & downspout|[units]\n",
    "Diversion": "Diversion (362)|1a|Clean water diversion|[units]\n",
    "Covered dry stack manure storage": "Waste Storage Facility (313)|3a|Dry stack pad|[units]\nRoof and Covers (367)|3b|Roof over waste storage|[units]\nRoof Runoff Structure (558)|3c|Gutter & downspout|[units]\nSubsurface Drain (606)|3d|Drain around storage|[units]\n",
    "Waste storage facility": "Waste Storage Facility (313)|3a|Waste storage facility|[units]\n",
    "Stacking pad": "Waste Storage Facility (313)|3a|Stacking pad|[units]\n",
    "Roof/gutter/outlet improvements": "Roof and Covers (367)|3a|Roof/cover|[units]\nRoof Runoff Structure (558)|3b|Gutter & downspout|[units]\nUnderground Outlet (620)|3c|Clean water outlet|[units]\n",
    "Prescribed grazing system": "Grazing Management (528)|2a|Prescribed grazing plan|[units]\nFence (382)|2b|Pasture fencing|[units]\nWatering Facility (614)|2c|Watering facility|[units]\nLivestock Pipeline (516)|2d|Water supply pipeline|[units]\n",
    "Fence and paddock divisions": "Fence (382)|2a|Pasture/paddock fencing|[units]\n",
    "Watering facility and livestock pipeline": "Watering Facility (614)|2a|Watering facility|[units]\nLivestock Pipeline (516)|2b|Water supply pipeline|[units]\n",
    "Trails and walkways / laneway": "Trails and Walkways (575)|2a|Livestock laneway|[units]\n",
    "Composting facility": "Animal Mortality Facility (316)|4a|Mortality composting facility|[units]\n",
}

BASIC_SHORT_FIELDS = [
    ("Farm Name", ""),
    ("Owner / Contact", ""),
    ("Address", ""),
    ("Town", ""),
    ("County", ""),
    ("State", "New York"),
    ("Operation Type", "beef cattle operation"),
    ("Watershed", ""),
    ("Basin", "Lake Ontario Basin"),
    ("Nearby Location / Reference", ""),
    ("Main Watercourse / Drainage", ""),
    ("Drainage / Topography Note", ""),
    ("Managed Acres", ""),
    ("Evaluation Date", ""),
    ("Attendees", ""),
    ("Housing / Animal Access", ""),
    ("Feed Program", ""),
    ("Manure Handling", ""),
    ("Application Timing", ""),
    ("Pasture Manure Note", ""),
]

# Each non-basic section gets short input boxes, then generated editable narratives.
SECTION_FACT_FIELDS = {
    "Barnyard": [
        ("Area name / field label", "barnyard / heavy use area"),
        ("Existing condition summary", ""),
        ("Area size", ""),
        ("Current use", "primary livestock transit point and feeding area"),
        ("Existing surface condition", "muddy, denuded, compacted, and lacking vegetative cover"),
        ("Animal traffic / feeding details", ""),
        ("Manure / bedding / TMP issue", ""),
        ("Runoff or water quality concern", "clean water is commingling with manure, resulting in sediment and nutrient-laden runoff"),
        ("Soil / drainage note", ""),
        ("Soil test / nutrient note", ""),
        ("Planned BMP", "covered heavy use area and restricted-use area"),
        ("Planned BMP dimensions", ""),
        ("Restricted / reseeded area", ""),
        ("Clean water handling", "gutters and underground outlets will convey clean water away from the production area"),
        ("Operational improvement", "manure can be scraped from the concrete area into the proposed manure storage"),
        ("Designer / agency", ""),
        ("Additional planned notes", ""),
    ],
    "Clean Stormwater": [
        ("Roof / clean water source", "barn roof runoff"),
        ("Existing condition summary", ""),
        ("Existing outlet / discharge location", ""),
        ("Area receiving clean water", ""),
        ("Current concern", "clean stormwater enters the livestock use area and becomes contaminated"),
        ("Working gutters / downspouts note", ""),
        ("Planned BMP", "roof runoff structure and underground outlet"),
        ("Planned outlet route", ""),
        ("Outlet daylight / discharge location", "stable, vegetated outlet location"),
        ("Length / size", ""),
        ("O&M note", "continue to operate and maintain gutters and downspouts in proper working condition"),
        ("Additional planned notes", ""),
    ],
    "Livestock Manure Storage": [
        ("Current manure storage practice", "temporary manure piles"),
        ("Existing condition summary", ""),
        ("Permanent storage present?", "no permanent manure storage facility is currently present on-site"),
        ("Barn cleanout frequency", ""),
        ("Bedding type", ""),
        ("Spreader / handling equipment", ""),
        ("Application / export note", ""),
        ("Observed manure concern", "manure-laden leachate was observed near the piles"),
        ("Nearby water / sensitive area", ""),
        ("Clean water contact note", "clean water from the farmstead encounters the manure piles"),
        ("Human waste separation note", "human sanitary waste is handled by a separate septic system"),
        ("Planned storage BMP", "covered dry stack manure storage"),
        ("PE / design note", "under the supervision of a Professional Engineer licensed in New York State"),
        ("AWM storage duration", "six (6) months"),
        ("AWM required capacity", ""),
        ("Storage dimensions", ""),
        ("Wall / stack height", ""),
        ("Emptying method", "bucketing manure into the spreader"),
        ("Roof / cover note", ""),
        ("Gutter / outlet / subsurface drain note", "gutters should be installed and attached to an underground outlet to divert clean water away from possible contamination sites; a subsurface drain may also be required"),
        ("Additional planned notes", ""),
    ],
    "Pasture": [
        ("Pasture acres / fields", ""),
        ("Existing condition summary", ""),
        ("Current grazing system", "continuous, year-round grazing pressure"),
        ("Problem area", ""),
        ("Existing pasture condition", "denuded, compacted, and lacking protective forage cover"),
        ("Resource concern", "exposed areas serve as a source of sediment and nutrient transport during precipitation events"),
        ("Water / stream / lane issue", ""),
        ("Planned BMP", "prescribed grazing plan"),
        ("Paddocks / divisions", ""),
        ("Water tanks", ""),
        ("Laneway / access plan", ""),
        ("Fence / gates / pipeline", ""),
        ("Management goal", "reduce the duration and intensity of livestock pressure in any one area and allow vegetation to recover"),
        ("Additional planned notes", ""),
    ],
    "Animal Mortalities": [
        ("Current mortality method", "burial"),
        ("Existing condition summary", ""),
        ("Burial depth / timeframe", "at least 6 feet deep within 72 hours"),
        ("Mortality site location", ""),
        ("Distance from watercourse", ""),
        ("Typical mortality rate", ""),
        ("Regulation note", "the farm follows New York State regulation"),
        ("Resource concern observed", "no significant resource concern was observed"),
        ("Planned / O&M action", "continue to maintain the burial site and surrounding area to minimize potential impacts to soil and water resources"),
    ],
    "Fuel / Petroleum": [
        ("Current petroleum storage", "no petroleum storage on farm"),
        ("Existing condition summary", ""),
        ("Storage condition / concern", "no resource concern was observed regarding petroleum storage"),
        ("Planned / O&M action", "if petroleum storage is added in the future, follow applicable DEC tank guidance and all local and state requirements"),
        ("Additional notes", ""),
    ],
    "Air Quality": [
        ("Manure type / storage", "solid manure with no liquid storage"),
        ("Existing condition summary", ""),
        ("Odor condition", "odor is minimal"),
        ("Dust condition", "dust is kept under control"),
        ("Ammonia / methane note", "there are no significant amounts of ammonia or methane"),
        ("Current concern", "no current air quality concerns are present"),
        ("Planned / O&M action", "follow the right rates and right placement of manure in the fields, and be mindful of nearby residents during windy conditions"),
        ("Additional notes", ""),
    ],
    "Driveways": [
        ("Current driveway condition", "driveways are generally maintained"),
        ("Existing condition summary", ""),
        ("Observed issue", "some manure was observed on the driveways from moving manure"),
        ("Resource concern", "manure on traffic areas can be tracked or moved by runoff if not cleaned up"),
        ("Planned / O&M action", "ensure driveways remain in good working order, add gravel as needed, and clean up manure from driveways"),
        ("Additional notes", ""),
    ],
    "Feed Storage": [
        ("Current feed storage", "hay is fed to livestock and no corn silage is stored on the farm"),
        ("Existing condition summary", ""),
        ("Storage location / method", ""),
        ("Observed concern", "no resource concerns for feed storage were observed"),
        ("Planned / O&M action", "continue operation and maintenance"),
        ("Additional notes", ""),
    ],
}

FIELD_HELP = {'Basic Info': {'Farm Name': 'Farm name exactly as it should appear in the report.', 'Owner / Contact': 'Primary farm owner/operator or contact for the evaluation.', 'Address': 'Main farmstead or production area address.', 'Town': 'Town or municipality where the farm is located.', 'County': 'County where the farm is located.', 'State': 'State for the farm location, usually New York.', 'Operation Type': 'Type of farm, such as beef cattle operation, dairy, horse farm, or mixed livestock.', 'Watershed': 'Watershed that receives drainage from the farmstead, if known.', 'Basin': 'Larger basin, such as Lake Ontario Basin or Lake Erie Basin.', 'Nearby Location / Reference': 'Short location reference, such as miles from a town, lake, or major road.', 'Main Watercourse / Drainage': 'Stream, ditch, creek, wetland, pond, or drainage feature on or near the farmstead.', 'Drainage / Topography Note': 'Short note about slope, low areas, surface drainage, wetness, or runoff direction.', 'Managed Acres': 'Total managed acres, crop acres, pasture acres, or acres tied to this evaluation.', 'Evaluation Date': 'Date the farmstead evaluation/site visit was completed.', 'Attendees': 'People present at the evaluation, including owner/operator and agency/planner names.', 'Housing / Animal Access': 'Where animals are housed and what areas they access, such as barn, pasture, barnyard, or loafing area.', 'Feed Program': 'What is fed, where it is fed, and whether feeding changes by season.', 'Manure Handling': 'How manure and bedding are handled, scraped, piled, stored, hauled, or spread.', 'Application Timing': 'When manure is normally spread or applied.', 'Pasture Manure Note': 'How manure deposited directly by animals in pasture/turnout areas is handled.', 'Animal Inventory': 'List each animal group on its own line, such as 20 mature cows, 1 bull, 15 calves.'}, 'Barnyard': {'Area name / field label': 'Name or label the area, such as barnyard, sacrifice lot, HUA, D1, or lot behind barn.', 'Area size': "Approximate size of the evaluated area, such as 0.9 acres or 130' x 110'.", 'Current use': 'How livestock use this area now: feeding, traveling, loafing, watering, barn access, etc.', 'Existing surface condition': 'Surface condition now: muddy, denuded, compacted, gravel, concrete, vegetated, saturated, rutted, etc.', 'Animal traffic / feeding details': 'Where animals concentrate and why, including feeder wagons, bale rings, waterers, gates, or travel routes.', 'Manure / bedding / TMP issue': 'Manure buildup, bedding, temporary manure piles, leachate, scraping, or cleanout concerns.', 'Runoff or water quality concern': 'How runoff could carry manure, sediment, or nutrients toward a ditch, stream, road ditch, wetland, or field.', 'Soil / drainage note': 'Soil type, wetness, seasonal saturation, drainage class, or low-area details if known.', 'Soil test / nutrient note': 'High soil test nutrients, pH, or comparison to other fields if useful.', 'Planned BMP': 'Planned practice/system, such as covered HUA, concrete pad, restricted-use area, fence, or seeding.', 'Planned BMP dimensions': "Size of the planned practice, such as 32' x 60' or 1,920 sq. ft.", 'Restricted / reseeded area': 'Area that will be fenced, reseeded, controlled, or changed to limited livestock access.', 'Clean water handling': 'How gutters, roof runoff, diversions, or underground outlets keep clean water away from manure.', 'Operational improvement': 'How the plan improves daily management, such as easier scraping, better footing, or less mud.', 'Designer / agency': 'Who designed or will design it, such as SWCD, NRCS, PE, or WNYCMA.', 'Additional planned notes': 'Any extra planned action detail that does not fit the other boxes.'}, 'Clean Stormwater': {'Roof / clean water source': 'Clean water source, such as barn roof runoff, gutters, downspouts, spring water, or upslope runoff.', 'Existing outlet / discharge location': 'Where clean water currently goes: barnyard, pasture, ditch, road ditch, HUA, manure area, etc.', 'Area receiving clean water': 'Area being saturated or affected, such as D1, feeding area, barnyard, or manure pile area.', 'Current concern': 'Why this clean water is a problem, usually because it contacts manure or creates mud/runoff.', 'Working gutters / downspouts note': 'Existing gutters/downspouts that work properly and should continue to be maintained.', 'Planned BMP': 'Planned clean water practice: roof runoff structure, underground outlet, diversion, or outlet stabilization.', 'Planned outlet route': 'Where the pipe, outlet, or diversion will carry clean water.', 'Outlet daylight / discharge location': 'Where the clean water will discharge, such as stable vegetated area or road ditch.', 'Length / size': 'Approximate lengths/sizes, such as 245 lf outlet or 60 lf gutter.', 'O&M note': 'Maintenance needed, such as cleaning gutters, checking outlets, and repairing erosion.', 'Additional planned notes': 'Any extra stormwater plan detail.'}, 'Livestock Manure Storage': {'Current manure storage practice': 'How manure is currently stored: temporary piles, daily spreading, bedded pack, stack pad, etc.', 'Permanent storage present?': 'Whether permanent manure storage exists and whether it is adequate.', 'Barn cleanout frequency': 'How often manure/bedding are removed from the barn or housing area.', 'Bedding type': 'Bedding material, such as straw, sawdust, wood shavings, sand, or none.', 'Spreader / handling equipment': 'Equipment used to move/spread manure, such as box spreader, skid steer, loader, or tractor.', 'Application / export note': 'Where manure goes, when it is spread, or whether any manure is exported.', 'Observed manure concern': 'Leachate, ponding, pile size, runoff, clean water contact, or manure near water.', 'Nearby water / sensitive area': 'Nearby ditches, streams, wetlands, road ditches, wells, flood-prone areas, or poorly drained soils.', 'Clean water contact note': 'Any clean water that reaches manure piles or storage areas.', 'Human waste separation note': 'Whether human sanitary waste is handled separately, usually by septic system.', 'Planned storage BMP': 'Planned storage practice, such as covered dry stack, waste storage facility, or stacking pad.', 'PE / design note': 'Engineering/design oversight if needed.', 'AWM storage duration': 'Planned storage period, such as six months, 180 days, or winter storage.', 'AWM required capacity': 'Required storage volume from AWM, such as 4,683 cubic feet.', 'Storage dimensions': "Planned storage dimensions, such as 32' x 34.6'.", 'Wall / stack height': 'Wall height or stack height, such as 4 foot walls with a 5 foot stack.', 'Emptying method': 'How storage will be emptied, such as bucketing into spreader or scraping with skid steer.', 'Roof / cover note': 'Roof or cover details for the storage.', 'Gutter / outlet / subsurface drain note': 'Clean water diversion, gutters, outlets, or subsurface drainage around storage.', 'Additional planned notes': 'Extra storage siting, soils, testing, or operation detail.'}, 'Pasture': {'Pasture acres / fields': 'Total pasture acres and field labels, such as 19.2 acres (A1-D1).', 'Current grazing system': 'Current use: continuous grazing, rotational grazing, seasonal turnout, year-round access, etc.', 'Problem area': 'Pasture area causing concern, such as around water tanks, gates, barn access, stream crossing, or D1.', 'Existing pasture condition': 'Cover/soil condition: overgrazed, denuded, compacted, adequate forage, bare around gates, etc.', 'Resource concern': 'Erosion, nutrient transport, sediment movement, poor livestock distribution, stream access, lack of recovery time, etc.', 'Water / stream / lane issue': 'Water tanks, pipelines, stream access, laneways, or animal travel paths affecting pasture management.', 'Planned BMP': 'Planned pasture BMP: prescribed grazing, fence, laneway, watering facility, pipeline, etc.', 'Paddocks / divisions': 'Planned pasture divisions or number of paddocks.', 'Water tanks': 'Planned water tanks or watering locations.', 'Laneway / access plan': 'Planned lane/access route and how cattle will move without damaging sensitive areas.', 'Fence / gates / pipeline': 'Planned fence, gates, pipeline, and related infrastructure.', 'Management goal': 'What the plan should accomplish: rest forage, improve distribution, reduce mud, protect water quality.', 'Additional planned notes': 'Extra pasture planning detail.'}, 'Animal Mortalities': {'Current mortality method': 'How mortalities are handled now: burial, composting, rendering, landfill, etc.', 'Burial depth / timeframe': 'Burial depth and timing if burial is used.', 'Mortality site location': 'Where the mortality site is located relative to barns, fields, roads, or property features.', 'Distance from watercourse': 'Approximate distance from streams, ditches, ponds, wetlands, wells, or watercourses.', 'Typical mortality rate': 'Normal number of mortalities per year, if known.', 'Regulation note': 'Whether the farm follows applicable NYS mortality disposal requirements.', 'Resource concern observed': 'Whether a concern was observed: none, poor location, shallow burial, runoff risk, scavenger issue, etc.', 'Planned / O&M action': 'Recommended action: continued compliance, monitoring, better site selection, or proper composting/burial management.'}, 'Fuel / Petroleum': {'Current petroleum storage': 'Tanks, containers, fuel type, or state that no petroleum storage is present.', 'Storage condition / concern': 'Condition, containment, leaks, spills, location, or whether no concern was observed.', 'Planned / O&M action': 'Maintenance or future requirements: DEC tank guidance, containment, spill cleanup, or monitoring.', 'Additional notes': 'Any other fuel/petroleum detail.'}, 'Air Quality': {'Manure type / storage': 'Manure type/storage: solid manure, liquid manure, bedded pack, covered stack, no liquid storage, etc.', 'Odor condition': 'Odor level or odor concern.', 'Dust condition': 'Dust from driveways, bedding, feeding, field work, or no issue.', 'Ammonia / methane note': 'Whether ammonia/methane is a concern, especially with liquid manure or enclosed storage.', 'Current concern': 'Whether any air quality concern is present.', 'Planned / O&M action': 'O&M actions to limit odor, dust, and air quality impacts.', 'Additional notes': 'Any other air quality detail.'}, 'Driveways': {'Current driveway condition': 'Driveway surface/condition: gravel, stable, rutted, muddy, eroding, maintained, etc.', 'Observed issue': 'Manure tracking, sediment, ruts, potholes, ponding, or no issue observed.', 'Resource concern': 'How driveway conditions could move sediment, manure, nutrients, or runoff offsite.', 'Planned / O&M action': 'Grading, gravel addition, cleanup, waterbars, ditching, or ongoing maintenance.', 'Additional notes': 'Any other driveway detail.'}, 'Feed Storage': {'Current feed storage': 'Feed stored on farm: hay, baleage, corn silage, grain, or no silage storage.', 'Storage location / method': 'Where/how feed is stored: barn, wrapped bales, bunker, bags, outside stacks, feeder wagon, etc.', 'Observed concern': 'Leachate, spoiled feed, runoff contact, wildlife, mud, or no concern observed.', 'Planned / O&M action': 'Cleanup, storage maintenance, runoff control, or continued O&M.', 'Additional notes': 'Any other feed storage detail.'}}


# Help text for the existing condition summary field used on each section tab.
for _sec in SECTION_FACT_FIELDS:
    FIELD_HELP.setdefault(_sec, {})["Existing condition summary"] = (
        "Briefly describe what is actually happening on the farm now. "
        "Use normal short notes; the program will clean the wording and build the narrative. "
        "Example: cattle travel across this concrete HUA to reach pasture, or temporary manure piles are located between the lower horse barn and riding arena."
    )

DEFAULT_TEXTS = {
    "Animal Inventory": "",
    "Introduction Narrative": "",
    "Overview Narrative": "",
    "Animal Housing and Feed Management Narrative": "",
    "Manure Management and Application Narrative": "",
}


WESTERN_NY_COUNTIES = {
    "allegany", "cattaraugus", "chautauqua", "erie", "genesee", "livingston",
    "monroe", "niagara", "ontario", "orleans", "seneca", "wayne", "wyoming", "yates"
}

CENTRAL_NY_COUNTIES = {
    "broome", "cayuga", "chenango", "cortland", "herkimer", "jefferson", "lewis",
    "madison", "oneida", "onondaga", "oswego", "otsego", "st. lawrence", "tioga", "tompkins"
}

def infer_ny_region(county, address=""):
    text = f"{county or ''} {address or ''}".lower()
    for c in WESTERN_NY_COUNTIES:
        if c in text:
            return "western New York"
    for c in CENTRAL_NY_COUNTIES:
        if c in text:
            return "central New York"
    return "New York"

def clean_watershed_phrase(watershed):
    """Clean common shorthand notes so the overview reads like a sentence."""
    text = (watershed or "").strip()
    if not text or text == "[watershed]":
        return text
    import re
    text = re.sub(r"\s+", " ", text)
    lower = text.lower()
    m = re.match(r"(?:the\s+)?farm(?:stead)?\s+spans\s+two\s+watersheds?\s+(.+)", lower, flags=re.I)
    if m:
        # Keep only the watershed names so the final sentence reads:
        # "The farmstead is within [name] Watershed of the [basin]."
        return text[m.start(1):].strip(" .")
    m = re.match(r"(?:the\s+)?farm(?:stead)?\s+is\s+within\s+(.+)", lower, flags=re.I)
    if m:
        return text[m.start(1):].strip(" .")
    return text.strip(" .")

def add_watershed_suffix(watershed):
    """Add the word Watershed unless it is already included."""
    text = (watershed or "").strip(" .")
    if not text or text == "[watershed]":
        return text
    import re
    if re.search(r"\bwatersheds?\b", text, flags=re.I):
        return text
    return f"{text} Watershed"

APP_BLUE = "#1f5f99"
APP_GREEN = "#2f7d32"
APP_BG = "#f3f7f2"
APP_PANEL = "#ffffff"
APP_PANEL_SOFT = "#fbfdfb"
APP_BORDER = "#d9e4d7"
APP_TEXT = "#223127"
APP_MUTED = "#5f6f66"
APP_ACCENT = "#2f7d32"
APP_ACCENT_DARK = "#245f26"
APP_BLUE_DARK = "#17466f"

class ScrollableFrame(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.canvas = tk.Canvas(self, highlightthickness=0, background=APP_BG)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.inner = ttk.Frame(self.canvas)
        self.inner.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas_window = self.canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.bind("<Configure>", lambda e: self.canvas.itemconfig(self.canvas_window, width=e.width))

        # Allow scrolling from anywhere inside the current work page, not just the scrollbar.
        # add="+" keeps all tabs registered; _is_descendant makes only the tab under the
        # mouse pointer scroll.
        self.bind_all("<MouseWheel>", self._on_mousewheel, add="+")
        self.bind_all("<Button-4>", self._on_mousewheel, add="+")
        self.bind_all("<Button-5>", self._on_mousewheel, add="+")

    def _is_descendant(self, widget):
        while widget is not None:
            if widget == self:
                return True
            widget = getattr(widget, "master", None)
        return False

    def _on_mousewheel(self, event):
        if not self._is_descendant(event.widget):
            return
        if getattr(event, "num", None) == 4:
            delta = -1
        elif getattr(event, "num", None) == 5:
            delta = 1
        else:
            delta = int(-1 * (event.delta / 120))
        self.canvas.yview_scroll(delta, "units")
        return "break"

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("WNY Crop Management - Farmstead Evaluation Generator")
        self.geometry("1280x900")
        self.minsize(1050, 720)
        self.configure(background=APP_BG)
        self.logo_img = None
        self.data_widgets = {}
        self.configure_styles()
        self.create_ui()

    def configure_styles(self):
        style = ttk.Style(self)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass

        # Modern visual styling only. Narrative/document logic is unchanged.
        style.configure("TFrame", background=APP_BG)
        style.configure("Header.TFrame", background=APP_PANEL)
        style.configure("Toolbar.TFrame", background=APP_PANEL)
        style.configure("Card.TFrame", background=APP_PANEL)

        style.configure("TLabel", background=APP_BG, foreground=APP_TEXT, font=("Segoe UI", 10))
        style.configure("HeaderTitle.TLabel", background=APP_PANEL, foreground=APP_BLUE_DARK, font=("Segoe UI", 22, "bold"))
        style.configure("HeaderSub.TLabel", background=APP_PANEL, foreground=APP_MUTED, font=("Segoe UI", 10))
        style.configure("ToolbarHint.TLabel", background=APP_PANEL, foreground=APP_MUTED, font=("Segoe UI", 9))
        style.configure("SectionTitle.TLabel", background=APP_BG, foreground=APP_ACCENT_DARK, font=("Segoe UI", 18, "bold"))
        style.configure("Help.TLabel", background=APP_BG, foreground=APP_MUTED, font=("Segoe UI", 9))
        style.configure("FieldLabel.TLabel", background=APP_BG, foreground=APP_TEXT, font=("Segoe UI", 10, "bold"))

        style.configure("TLabelframe", background=APP_BG, bordercolor=APP_BORDER, relief="solid")
        style.configure("TLabelframe.Label", background=APP_BG, foreground=APP_BLUE_DARK, font=("Segoe UI", 10, "bold"))
        style.configure("Card.TLabelframe", background=APP_BG, bordercolor=APP_BORDER, relief="solid")
        style.configure("Card.TLabelframe.Label", background=APP_BG, foreground=APP_ACCENT_DARK, font=("Segoe UI", 11, "bold"))

        style.configure("TNotebook", background=APP_BG, borderwidth=0, tabmargins=(4, 6, 4, 0))
        style.configure("TNotebook.Tab", padding=(16, 9), font=("Segoe UI", 9, "bold"), background="#e8efe6", foreground=APP_MUTED)
        style.map("TNotebook.Tab",
                  background=[("selected", APP_PANEL), ("active", "#f0f6ee")],
                  foreground=[("selected", APP_ACCENT_DARK), ("active", APP_BLUE_DARK)])

        style.configure("TEntry", fieldbackground=APP_PANEL, bordercolor=APP_BORDER, lightcolor=APP_BORDER, darkcolor=APP_BORDER, padding=4)
        style.configure("TCombobox", fieldbackground=APP_PANEL, bordercolor=APP_BORDER, lightcolor=APP_BORDER, darkcolor=APP_BORDER, padding=4)

        style.configure("TButton", font=("Segoe UI", 9), padding=(10, 6), background="#eef4ec", foreground=APP_TEXT, bordercolor=APP_BORDER)
        style.map("TButton", background=[("active", "#dfeade")])
        style.configure("Accent.TButton", font=("Segoe UI", 10, "bold"), padding=(12, 7), background=APP_ACCENT, foreground="white", bordercolor=APP_ACCENT_DARK)
        style.map("Accent.TButton", background=[("active", APP_ACCENT_DARK)], foreground=[("active", "white")])
        style.configure("Primary.TButton", font=("Segoe UI", 10, "bold"), padding=(12, 7), background=APP_BLUE, foreground="white", bordercolor=APP_BLUE_DARK)
        style.map("Primary.TButton", background=[("active", APP_BLUE_DARK)], foreground=[("active", "white")])
        style.configure("Subtle.TButton", font=("Segoe UI", 9), padding=(10, 6), background="#eef4ec", foreground=APP_MUTED, bordercolor=APP_BORDER)

    def load_logo(self):
        logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "wny_crop_management_logo.png")
        if os.path.exists(logo_path):
            try:
                img = tk.PhotoImage(file=logo_path)
                if img.width() > 560:
                    img = img.subsample(2, 2)
                self.logo_img = img
            except tk.TclError:
                self.logo_img = None

    def create_ui(self):
        self.load_logo()

        header = ttk.Frame(self, padding=(18, 14), style="Header.TFrame")
        header.pack(fill="x")
        if self.logo_img:
            ttk.Label(header, image=self.logo_img, background=APP_PANEL).pack(side="left", padx=(0, 20))
        title_area = ttk.Frame(header, style="Header.TFrame")
        title_area.pack(side="left", fill="x", expand=True)
        ttk.Label(title_area, text="Farmstead Evaluation Generator", style="HeaderTitle.TLabel").pack(anchor="w")
        ttk.Label(title_area, text="Quick Build Wizard • reviewed-style narratives • Word document output", style="HeaderSub.TLabel").pack(anchor="w", pady=(4, 0))
        ttk.Label(title_area, text="Western New York Crop Management Association Cooperative", style="HeaderSub.TLabel").pack(anchor="w", pady=(2, 0))

        top_outer = ttk.Frame(self, padding=(12, 10, 12, 8))
        top_outer.pack(fill="x")
        top = ttk.Frame(top_outer, padding=(10, 10), style="Toolbar.TFrame")
        top.pack(fill="x")
        ttk.Button(top, text="Save Input", command=self.save_json, style="Subtle.TButton").pack(side="left", padx=4)
        ttk.Button(top, text="Open Input", command=self.load_json, style="Subtle.TButton").pack(side="left", padx=4)
        ttk.Button(top, text="Quick Build Wizard", command=self.open_quick_build_wizard, style="Primary.TButton").pack(side="left", padx=(12, 4))
        ttk.Button(top, text="Refresh Narratives", command=self.refresh_all_narratives, style="Accent.TButton").pack(side="left", padx=4)
        ttk.Button(top, text="Generate Word Document", command=self.generate_docx, style="Accent.TButton").pack(side="left", padx=4)
        ttk.Button(top, text="Clear Section", command=self.clear_current_section, style="Subtle.TButton").pack(side="left", padx=(12, 4))
        ttk.Button(top, text="Clear All", command=self.clear_all_text_boxes, style="Subtle.TButton").pack(side="left", padx=4)
        ttk.Label(top, text="Review generated narratives before finalizing.", style="ToolbarHint.TLabel").pack(side="right", padx=8)

        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True, padx=12, pady=(2, 12))
        self.tab_names = []
        for idx, section in enumerate(SECTIONS, start=1):
            tab = ScrollableFrame(self.notebook)
            self.notebook.add(tab, text=f"{idx}. {section}")
            self.tab_names.append(section)
            if section == "Basic Info":
                self.build_basic_info(tab.inner, section)
                qtab = ScrollableFrame(self.notebook)
                self.notebook.add(qtab, text="Quick Build Wizard")
                self.quick_build_tab_index = len(self.tab_names)
                self.tab_names.append("Quick Build Wizard")
                self.build_quick_build_wizard(qtab.inner)
            else:
                self.build_section_tab(tab.inner, section)

    def add_entry(self, parent, label, default="", help_text=""):
        frame = ttk.Frame(parent)
        frame.pack(fill="x", pady=4, padx=(0, 10))
        row = ttk.Frame(frame)
        row.pack(fill="x")
        ttk.Label(row, text=label, width=34, style="FieldLabel.TLabel").pack(side="left")
        entry = ttk.Entry(row)
        entry.pack(side="left", fill="x", expand=True)
        entry.insert(0, default)
        if help_text:
            ttk.Label(
                frame,
                text="What to enter: " + help_text,
                style="Help.TLabel",
                wraplength=1000,
                justify="left"
            ).pack(anchor="w", padx=(240, 0), pady=(1, 3))
        return entry

    def add_text(self, parent, label, height=5, default=""):
        ttk.Label(parent, text=label, font=("Segoe UI", 10, "bold"), foreground=APP_BLUE_DARK).pack(anchor="w", pady=(12, 3))
        text = tk.Text(
            parent,
            height=height,
            wrap="word",
            undo=True,
            font=("Segoe UI", 10),
            relief="solid",
            borderwidth=1,
            padx=8,
            pady=6,
            background=APP_PANEL_SOFT,
            foreground=APP_TEXT,
            insertbackground=APP_TEXT,
            highlightthickness=1,
            highlightbackground=APP_BORDER,
            highlightcolor=APP_ACCENT,
        )
        text.pack(fill="x", expand=False, padx=(0, 10))
        text.insert("1.0", default)
        return text

    def build_basic_info(self, parent, section):
        ttk.Label(parent, text="Basic Info", style="SectionTitle.TLabel").pack(anchor="w", pady=(4, 8))
        ttk.Label(parent, text="Only type short farm-specific facts here. The program writes the front-page narratives for you.").pack(anchor="w", pady=(0, 8))
        self.data_widgets[section] = {}

        quick = ttk.LabelFrame(parent, text="Short inputs used to write the front page narrative", padding=12, style="Card.TLabelframe")
        quick.pack(fill="x", padx=(0, 8), pady=(0, 8))
        for field, default in BASIC_SHORT_FIELDS:
            w = self.add_entry(quick, field, default, FIELD_HELP.get(section, {}).get(field, ""))
            self.data_widgets[section][field] = w

        ttk.Label(
            quick,
            text="What to enter: " + FIELD_HELP.get(section, {}).get("Animal Inventory", ""),
            foreground="#555555",
            wraplength=1000,
            justify="left"
        ).pack(anchor="w", pady=(8, 0))
        inv = self.add_text(quick, "Animal Inventory - one line per animal group", height=5, default=DEFAULT_TEXTS["Animal Inventory"])
        self.data_widgets[section]["Animal Inventory"] = inv

        help_text = (
            "Examples for short inputs:\n"
            "Nearby Location / Reference: approximately 12 miles north of Lockport and 6 miles south of Lake Ontario\n"
            "Main Watercourse / Drainage: East Branch of 12-Mile Creek flows through the property\n"
            "Drainage / Topography Note: minimal surface drainage due to local topography\n"
            "Housing / Animal Access: cattle have continuous access to a 19.2-acre pasture (A1, B1, C1, D1), which encompasses the main barn\n"
            "Feed Program: hay is provided in the covered main barn during winter and inclement weather, with a 4-5 bale feeder wagon placed in pasture when cattle are outside\n"
            "Manure Handling: manure and bedding from the barn are managed through uncovered temporary manure piles located east of the main barn\n"
            "Application Timing: land application occurs from spring through fall when weather and field conditions are suitable\n"
            "Pasture Manure Note: within the pasture area, manure is deposited naturally by the herd year-round"
        )
        examples = ttk.LabelFrame(parent, text="Input examples", padding=12, style="Card.TLabelframe")
        examples.pack(fill="x", padx=(0, 8), pady=(8, 8))
        ttk.Label(examples, text=help_text, style="Help.TLabel", wraplength=1000, justify="left").pack(anchor="w")

        generated = ttk.LabelFrame(parent, text="Generated front page narratives - editable before creating Word document", padding=12, style="Card.TLabelframe")
        generated.pack(fill="x", padx=(0, 8), pady=(8, 8))
        ttk.Button(generated, text="Generate / Refresh Front Page Narratives", command=self.refresh_front_page_narratives).pack(anchor="w", pady=(0, 6))
        for field in [
            "Introduction Narrative",
            "Overview Narrative",
            "Animal Housing and Feed Management Narrative",
            "Manure Management and Application Narrative",
        ]:
            h = 6 if field in ["Introduction Narrative", "Overview Narrative"] else 5
            w = self.add_text(generated, field, height=h, default=DEFAULT_TEXTS.get(field, ""))
            self.data_widgets[section][field] = w

    def build_section_tab(self, parent, section):
        ttk.Label(parent, text=section, style="SectionTitle.TLabel").pack(anchor="w", pady=(4, 8))
        ttk.Label(parent, text="Enter short facts. Then generate this section's Existing Conditions and Planned Actions narratives.").pack(anchor="w", pady=(0,8))
        self.data_widgets[section] = {}

        facts = ttk.LabelFrame(parent, text="Short inputs for this section", padding=12, style="Card.TLabelframe")
        facts.pack(fill="x", padx=(0, 8), pady=(0, 8))
        for field, default in SECTION_FACT_FIELDS.get(section, []):
            w = self.add_entry(facts, field, default, FIELD_HELP.get(section, {}).get(field, ""))
            self.data_widgets[section][field] = w

        btns = ttk.Frame(parent)
        btns.pack(fill="x", padx=(0,8), pady=(4, 2))
        ttk.Button(btns, text=f"Generate / Refresh {section} Narratives", command=lambda s=section: self.refresh_section_narratives(s, show_msg=True)).pack(side="left", padx=(0, 6))
        ttk.Button(btns, text="Use No Resource Concern / O&M Narrative", command=lambda s=section: self.use_no_resource_concern_narrative(s)).pack(side="left", padx=(0, 6))

        generated = ttk.LabelFrame(parent, text="Generated narratives - editable", padding=12, style="Card.TLabelframe")
        generated.pack(fill="x", padx=(0, 8), pady=(8, 8))
        ex = self.add_text(generated, "Existing Conditions narrative", height=11)
        self.data_widgets[section]["Existing Conditions narrative"] = ex
        pl = self.add_text(generated, "Planned Actions narrative", height=11)
        self.data_widgets[section]["Planned Actions narrative"] = pl
        ttk.Label(
            generated,
            text="What to enter: List each planned NRCS practice on its own line using this format: Standard | Item # | Description | Units. Example: Heavy Use Area Protection (561)|1d|Concrete pad with curbs|1,920 Sq. ft. Leave only the header if no standard is needed.",
            foreground="#555555",
            wraplength=1000,
            justify="left"
        ).pack(anchor="w", pady=(8, 0))
        st = self.add_text(generated, "NRCS Standards and Units Used", height=8, default="Standard | Item # | Description | Units\n")
        self.data_widgets[section]["NRCS Standards and Units Used"] = st

    def add_combo(self, parent, label, values, default="", help_text=""):
        frame = ttk.Frame(parent)
        frame.pack(fill="x", pady=4, padx=(0, 10))
        row = ttk.Frame(frame)
        row.pack(fill="x")
        ttk.Label(row, text=label, width=34, style="FieldLabel.TLabel").pack(side="left")
        combo = ttk.Combobox(row, values=values, state="readonly")
        combo.pack(side="left", fill="x", expand=True)
        combo.set(default or (values[0] if values else ""))
        if help_text:
            ttk.Label(frame, text="What to enter: " + help_text, style="Help.TLabel", wraplength=1000, justify="left").pack(anchor="w", padx=(240, 0), pady=(1, 3))
        return combo

    def add_multi_condition_selector(self, parent, label, values, default="", help_text=""):
        """Searchable checklist for choosing one or more existing condition statements.

        This replaces the old Ctrl-click listbox with normal checkboxes, a filter box,
        and Clear/Select Visible controls so first-time users can make multiple selections
        without needing to know keyboard shortcuts.
        """
        frame = ttk.Frame(parent)
        frame.pack(fill="x", pady=4, padx=(0, 10))

        header = ttk.Frame(frame)
        header.pack(fill="x")
        ttk.Label(header, text=label, width=34).pack(side="left", anchor="n")

        selector = ttk.Frame(header)
        selector.pack(side="left", fill="both", expand=True)

        search_row = ttk.Frame(selector)
        search_row.pack(fill="x", pady=(0, 3))
        search_var = tk.StringVar(value="")
        search_entry = ttk.Entry(search_row, textvariable=search_var)
        search_entry.pack(side="left", fill="x", expand=True)
        search_entry.insert(0, "Search existing conditions...")
        search_entry.configure(foreground="#777777")

        btn_row = ttk.Frame(search_row)
        btn_row.pack(side="left", padx=(6, 0))

        canvas_wrap = ttk.Frame(selector)
        canvas_wrap.pack(fill="both", expand=True)
        canvas = tk.Canvas(canvas_wrap, height=155, highlightthickness=1, highlightbackground="#d0d0d0")
        scrollbar = ttk.Scrollbar(canvas_wrap, orient="vertical", command=canvas.yview)
        check_frame = ttk.Frame(canvas)
        window_id = canvas.create_window((0, 0), window=check_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        vars_by_value = {value: tk.BooleanVar(value=False) for value in values}
        visible_values = []
        checkbuttons = []

        def _on_frame_configure(event=None):
            canvas.configure(scrollregion=canvas.bbox("all"))

        def _on_canvas_configure(event):
            canvas.itemconfigure(window_id, width=event.width)

        check_frame.bind("<Configure>", _on_frame_configure)
        canvas.bind("<Configure>", _on_canvas_configure)

        def _clear_placeholder(event=None):
            if search_entry.get() == "Search existing conditions...":
                search_entry.delete(0, "end")
                search_entry.configure(foreground="#000000")

        def _restore_placeholder(event=None):
            if not search_entry.get().strip():
                search_entry.insert(0, "Search existing conditions...")
                search_entry.configure(foreground="#777777")
                _refresh_options()

        def _current_search_text():
            text = search_var.get().strip()
            return "" if text == "Search existing conditions..." else text.lower()

        def _refresh_options(*args):
            nonlocal visible_values, checkbuttons
            for cb in checkbuttons:
                cb.destroy()
            checkbuttons = []
            visible_values = []
            search_text = _current_search_text()
            matches = [v for v in values if not search_text or search_text in v.lower()]
            for value in matches:
                cb = ttk.Checkbutton(check_frame, text=value, variable=vars_by_value[value])
                cb.pack(anchor="w", fill="x", padx=4, pady=1)
                checkbuttons.append(cb)
                visible_values.append(value)
            if not matches:
                lbl = ttk.Label(check_frame, text="No matching conditions found.", foreground="#777777")
                lbl.pack(anchor="w", padx=4, pady=4)
                checkbuttons.append(lbl)
            _on_frame_configure()
            canvas.yview_moveto(0)

        def _clear_selected():
            for var in vars_by_value.values():
                var.set(False)

        def _select_visible():
            for value in visible_values:
                vars_by_value[value].set(True)

        ttk.Button(btn_row, text="Select Visible", command=_select_visible).pack(side="left", padx=(0, 3))
        ttk.Button(btn_row, text="Clear", command=_clear_selected).pack(side="left")

        search_entry.bind("<FocusIn>", _clear_placeholder)
        search_entry.bind("<FocusOut>", _restore_placeholder)
        search_var.trace_add("write", _refresh_options)

        def _on_mousewheel(event):
            delta = -1 * int(event.delta / 120) if event.delta else 0
            canvas.yview_scroll(delta, "units")

        canvas.bind("<MouseWheel>", _on_mousewheel)
        check_frame.bind("<MouseWheel>", _on_mousewheel)

        _refresh_options()

        if default:
            parts = [p.strip() for p in (default or "").replace("\n", ";").split(";") if p.strip()]
            for option in values:
                if option in parts:
                    vars_by_value[option].set(True)

        if help_text:
            ttk.Label(
                frame,
                text="What to enter: " + help_text + " Use the search box to narrow the list, then check every condition that applies. You can choose multiple items without holding Ctrl.",
                style="Help.TLabel",
                wraplength=1000,
                justify="left"
            ).pack(anchor="w", padx=(240, 0), pady=(1, 3))

        class MultiConditionSelection:
            def __init__(self, all_values, variables):
                self.values = all_values
                self.vars_by_value = variables
            def get_value(self):
                selected = [value for value in self.values if self.vars_by_value[value].get()]
                return "; ".join([x for x in selected if x])
            def set_value(self, value):
                for var in self.vars_by_value.values():
                    var.set(False)
                parts = [p.strip() for p in (value or "").replace("\n", ";").split(";") if p.strip()]
                if not parts:
                    return
                for option in self.values:
                    if option in parts:
                        self.vars_by_value[option].set(True)

        return MultiConditionSelection(values, vars_by_value)

    def add_multi_practice_dropdown(self, parent, label, choices, default="", help_text=""):
        """Searchable Planned BMP/action selector with an editable practice/dimensions table.

        This version uses a normal Entry + Listbox instead of forcing open a Combobox.
        That avoids the Windows/Tk crash that happened when clicking into the search box.
        """
        frame = ttk.Frame(parent)
        frame.pack(fill="x", pady=4, padx=(0, 10))

        row = ttk.Frame(frame)
        row.pack(fill="x")
        ttk.Label(row, text=label, width=34, style="FieldLabel.TLabel").pack(side="left")

        search_var = tk.StringVar(value="")
        search_entry = ttk.Entry(row, textvariable=search_var)
        search_entry.pack(side="left", fill="x", expand=True)

        if help_text:
            ttk.Label(
                frame,
                text="What to enter: " + help_text + " Start typing a practice name or standard number. Matching standards will show below. Select one and click Add Practice to Table. After adding a practice, fill in the Item # and Units/Dimensions directly in the table next to that practice.",
                style="Help.TLabel",
                wraplength=1000,
                justify="left"
            ).pack(anchor="w", padx=(240, 0), pady=(1, 3))

        suggestion_frame = ttk.Frame(frame)
        suggestion_frame.pack(fill="x", padx=(240, 0), pady=(0, 4))
        suggestion_list = tk.Listbox(suggestion_frame, height=5, exportselection=False)
        suggestion_scroll = ttk.Scrollbar(suggestion_frame, orient="vertical", command=suggestion_list.yview)
        suggestion_list.configure(yscrollcommand=suggestion_scroll.set)
        suggestion_list.pack(side="left", fill="x", expand=True)
        suggestion_scroll.pack(side="right", fill="y")

        table_frame = ttk.Frame(frame)
        table_frame.pack(fill="x", padx=(240, 0), pady=(0, 4))

        class PracticeTable:
            def __init__(self, outer, table_parent, practice_choices, entry_widget, list_widget):
                self.outer = outer
                self.parent = table_parent
                self.choices = practice_choices
                self.entry = entry_widget
                self.listbox = list_widget
                self.rows = []
                headers = ["Planned Practice", "Item #", "Description", "Units / Dimensions", ""]
                widths = [34, 10, 28, 24, 8]
                for col, (h, w) in enumerate(zip(headers, widths)):
                    ttk.Label(self.parent, text=h, font=("Segoe UI", 9, "bold"), width=w).grid(row=0, column=col, sticky="w", padx=2, pady=(0, 2))
                for col in range(4):
                    self.parent.grid_columnconfigure(col, weight=1 if col in (0, 2, 3) else 0)

            def _practice_parts(self, value):
                value = (value or "").strip()
                if value.endswith(")") and "(" in value:
                    desc = value.rsplit("(", 1)[0].strip()
                    code = value.rsplit("(", 1)[1].rstrip(")").strip()
                    if code.isdigit():
                        return f"{desc} ({code})", desc
                return value, value

            def _selected_or_typed_value(self):
                selection = self.listbox.curselection()
                if selection:
                    return self.listbox.get(selection[0]).strip()
                typed = self.entry.get().strip()
                matches = [c for c in self.choices if typed.lower() in c.lower()] if typed else []
                if len(matches) == 1:
                    return matches[0]
                return typed

            def add(self, value=None, item="", desc=None, units=""):
                value = (value or self._selected_or_typed_value()).strip()
                if not value:
                    return
                if value == "No structural BMP - continue O&M":
                    self.clear()
                    return
                # If a BMP system name is selected, add all component practices for that system.
                if value in BMP_SYSTEM_COMPONENTS:
                    for practice, code in BMP_SYSTEM_COMPONENTS[value]:
                        self.add(f"{practice} ({code})")
                    return
                practice, default_desc = self._practice_parts(value)
                existing_practices = [r["practice"].get().strip() for r in self.rows]
                if practice in existing_practices:
                    return
                r_index = len(self.rows) + 1
                practice_entry = ttk.Entry(self.parent, width=34)
                practice_entry.insert(0, practice)
                practice_entry.grid(row=r_index, column=0, sticky="ew", padx=2, pady=2)
                item_entry = ttk.Entry(self.parent, width=10)
                item_entry.insert(0, item)
                item_entry.grid(row=r_index, column=1, sticky="ew", padx=2, pady=2)
                desc_entry = ttk.Entry(self.parent, width=28)
                desc_entry.insert(0, default_desc if desc is None else desc)
                desc_entry.grid(row=r_index, column=2, sticky="ew", padx=2, pady=2)
                units_entry = ttk.Entry(self.parent, width=24)
                units_entry.insert(0, units)
                units_entry.grid(row=r_index, column=3, sticky="ew", padx=2, pady=2)
                row_obj = {"practice": practice_entry, "item": item_entry, "desc": desc_entry, "units": units_entry}
                def remove_row(ro=row_obj):
                    for w in ro.values():
                        w.destroy()
                    self.rows.remove(ro)
                    self.regrid()
                remove_btn = ttk.Button(self.parent, text="Remove", command=remove_row)
                remove_btn.grid(row=r_index, column=4, sticky="ew", padx=2, pady=2)
                row_obj["remove"] = remove_btn
                self.rows.append(row_obj)
                self.entry.delete(0, tk.END)
                refresh_suggestions()

            def regrid(self):
                for idx, ro in enumerate(self.rows, start=1):
                    ro["practice"].grid(row=idx, column=0, sticky="ew", padx=2, pady=2)
                    ro["item"].grid(row=idx, column=1, sticky="ew", padx=2, pady=2)
                    ro["desc"].grid(row=idx, column=2, sticky="ew", padx=2, pady=2)
                    ro["units"].grid(row=idx, column=3, sticky="ew", padx=2, pady=2)
                    ro["remove"].grid(row=idx, column=4, sticky="ew", padx=2, pady=2)

            def clear(self):
                for ro in list(self.rows):
                    for w in ro.values():
                        w.destroy()
                self.rows = []

            def get_value(self):
                lines = []
                for ro in self.rows:
                    practice = ro["practice"].get().strip()
                    if not practice:
                        continue
                    item = ro["item"].get().strip() or "[item #]"
                    desc = ro["desc"].get().strip() or _practice_name_from_line(practice).rsplit("(", 1)[0].strip()
                    units = ro["units"].get().strip() or "[units]"
                    lines.append(f"{practice}|{item}|{desc}|{units}")
                return "\n".join(lines)

            def set_value(self, value):
                self.clear()
                for line in (value or "").splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    if "|" in line:
                        parts = [p.strip() for p in line.split("|")]
                        parts += [""] * (4 - len(parts))
                        self.add(parts[0], parts[1], parts[2], parts[3])
                    else:
                        self.add(line)

        def refresh_suggestions(event=None):
            typed = search_var.get().lower().strip()
            filtered = [c for c in choices if typed in c.lower()] if typed else choices[:]
            suggestion_list.delete(0, tk.END)
            for c in filtered[:100]:
                suggestion_list.insert(tk.END, c)

        practice_table = PracticeTable(self, table_frame, choices, search_entry, suggestion_list)

        button_row = ttk.Frame(frame)
        button_row.pack(fill="x", padx=(240, 0), pady=(0, 2))
        ttk.Button(button_row, text="Add Practice to Table", command=lambda: practice_table.add()).pack(side="left", padx=(0, 6))
        ttk.Button(button_row, text="Clear Practice Table", command=practice_table.clear).pack(side="left")

        search_entry.bind("<KeyRelease>", refresh_suggestions)
        search_entry.bind("<FocusIn>", refresh_suggestions)
        suggestion_list.bind("<Double-Button-1>", lambda _event: practice_table.add())
        suggestion_list.bind("<Return>", lambda _event: practice_table.add())
        refresh_suggestions()
        return practice_table

    def add_practice_search_box(self, parent, section_name):
        frame = ttk.LabelFrame(parent, text="Add multiple planned BMP practices", padding=12, style="Card.TLabelframe")
        frame.pack(fill="x", padx=(0, 8), pady=(6, 6))
        ttk.Label(
            frame,
            text="What to enter: Start typing a practice name or standard number, select the matching practice, then click Add Selected Practice. Add as many planned practices as needed for this section.",
            style="Help.TLabel",
            wraplength=1000,
            justify="left"
        ).pack(anchor="w", pady=(0, 4))

        search_row = ttk.Frame(frame)
        search_row.pack(fill="x", pady=(0, 4))
        ttk.Label(search_row, text="Search standards", width=34).pack(side="left")
        search_entry = ttk.Entry(search_row)
        search_entry.pack(side="left", fill="x", expand=True)

        listbox = tk.Listbox(frame, height=5, selectmode="browse", exportselection=False)
        listbox.pack(fill="x", padx=(240, 0), pady=(0, 4))

        selected_text = tk.Text(frame, height=4, wrap="word")
        selected_text.pack(fill="x", padx=(240, 0), pady=(0, 4))

        def refresh_list(*_args):
            q = search_entry.get().lower().strip()
            listbox.delete(0, "end")
            matches = []
            for item in ALL_BMP_CHOICES:
                if not q or q in item.lower():
                    matches.append(item)
            for item in matches[:80]:
                listbox.insert("end", item)

        def add_selected():
            selection = listbox.curselection()
            if not selection:
                typed = search_entry.get().strip()
                if typed:
                    candidates = [x for x in ALL_BMP_CHOICES if typed.lower() in x.lower()]
                    if candidates:
                        value = candidates[0]
                    else:
                        value = typed
                else:
                    return
            else:
                value = listbox.get(selection[0])
            existing = [ln.strip() for ln in selected_text.get("1.0", "end").splitlines() if ln.strip()]
            if value not in existing:
                selected_text.insert("end", value + "\n")

        def clear_selected():
            selected_text.delete("1.0", "end")

        button_row = ttk.Frame(frame)
        button_row.pack(fill="x", padx=(240, 0), pady=(0, 2))
        ttk.Button(button_row, text="Add Selected Practice", command=add_selected).pack(side="left", padx=(0, 6))
        ttk.Button(button_row, text="Clear Selected Practices", command=clear_selected).pack(side="left")

        search_entry.bind("<KeyRelease>", refresh_list)
        listbox.bind("<Double-Button-1>", lambda _event: add_selected())
        refresh_list()
        return selected_text

    def open_quick_build_wizard(self):
        if hasattr(self, "quick_build_tab_index"):
            self.notebook.select(self.quick_build_tab_index)

    def build_quick_build_wizard(self, parent):
        section = "Quick Build Wizard"
        self.data_widgets[section] = {}
        ttk.Label(parent, text="Quick Build Wizard", style="SectionTitle.TLabel").pack(anchor="w", pady=(4, 8))
        ttk.Label(
            parent,
            text="Use this page to build the farmstead evaluation faster. The narrative is now written primarily from the Brief current existing condition box. The program uses that text, the resource-concern yes/no selection, the planned practices table, and the location/size fields to create a cleaner reviewed-style narrative.",
            wraplength=1100,
            justify="left"
        ).pack(anchor="w", pady=(0, 10))
        ttk.Button(parent, text="Build ALL Sections from Quick Wizard", command=self.build_all_from_quick_wizard, style="Accent.TButton").pack(anchor="w", pady=(0, 12))

        for sec in QUICK_BUILD_SECTIONS:
            box = ttk.LabelFrame(parent, text=sec, padding=12, style="Card.TLabelframe")
            box.pack(fill="x", padx=(0, 8), pady=(0, 10))
            concern = self.add_combo(
                box,
                f"{sec} has resource concern?",
                ["No", "Yes"],
                "No",
                "Select Yes if the existing condition is causing or could cause a soil, water, manure, leachate, runoff, odor, access, or management concern. Select No if the section only needs a short O&M narrative."
            )
            self.data_widgets[section][f"{sec} - Concern"] = concern

            existing_summary = self.add_text(
                box,
                "Brief current existing condition",
                height=5,
                default=""
            )
            ttk.Label(
                box,
                text="What to enter: Briefly describe what is currently happening. Write this like field notes, not a polished paragraph. Include the existing practice, what was observed, how the area is used, and why it may or may not be a concern. The program will clean punctuation, spacing, capitalization, and build the NRCS-style narrative from this box.",
                style="Help.TLabel",
                wraplength=1000,
                justify="left"
            ).pack(anchor="w", padx=(240, 0), pady=(1, 6))
            self.data_widgets[section][f"{sec} - Existing Summary"] = existing_summary

            selected_practices = self.add_multi_practice_dropdown(
                box,
                "Planned BMP / action",
                all_bmp_dropdown_choices(),
                "No structural BMP - continue O&M",
                "Start typing a practice name or standard number, choose the matching practice, then click Add Practice. You can add multiple practices for the same section. Fill in Item #, Description, and Units/Dimensions in the table."
            )
            self.data_widgets[section][f"{sec} - Selected Practices"] = selected_practices
            self.data_widgets[section][f"{sec} - BMP"] = selected_practices

            size = self.add_entry(box, "Size / units / dimensions", "", "Optional general size information for the evaluated area or planned BMP. Practice-specific sizes should still go in the Planned BMP table.")
            self.data_widgets[section][f"{sec} - Size"] = size
            condition_name = self.add_entry(box, "Existing condition name", "", "Enter what this existing condition should be called in the narrative, such as barnyard, sacrifice lot, temporary manure pile area, east barn roof runoff, pasture #3, driveway, or feed storage area.")
            self.data_widgets[section][f"{sec} - Condition Name"] = condition_name
            loc = self.add_entry(box, "Area / BMP location", "", "Enter where the evaluated area or planned BMP is located, such as field ID, barn side, pasture name, road ditch, lot name, or nearby watercourse.")
            self.data_widgets[section][f"{sec} - Location"] = loc
            details = self.add_entry(box, "Other key details", "", "Optional: add one extra detail that must appear in the narrative, such as soil/wetland proximity, equipment, cleanout frequency, animal numbers, or design note.")
            self.data_widgets[section][f"{sec} - Details"] = details
            std = self.add_text(box, "Optional standards override", height=4, default="")
            self.data_widgets[section][f"{sec} - Standards"] = std

        ttk.Button(parent, text="Build ALL Sections from Quick Wizard", command=self.build_all_from_quick_wizard, style="Accent.TButton").pack(anchor="w", pady=(8, 16))

    def build_all_from_quick_wizard(self):
        for sec in QUICK_BUILD_SECTIONS:
            self.apply_quick_build_section(sec)
        messagebox.showinfo("Quick Build complete", "The section tabs were filled from the Quick Build Wizard. Review/edit the detailed section tabs before creating the Word document.")

    def apply_quick_build_section(self, section):
        qw = "Quick Build Wizard"
        concern = self.get_widget_value(qw, f"{section} - Concern")
        selected_practices = self.get_widget_value(qw, f"{section} - Selected Practices")
        bmp = practice_lines_to_sentence(selected_practices)
        size = self.get_widget_value(qw, f"{section} - Size")
        existing_summary = self.get_widget_value(qw, f"{section} - Existing Summary")
        condition_name = self.get_widget_value(qw, f"{section} - Condition Name")
        loc = self.get_widget_value(qw, f"{section} - Location")
        details = self.get_widget_value(qw, f"{section} - Details")
        standards_override = self.get_widget_value(qw, f"{section} - Standards")

        def set_if(field, value):
            if field in self.data_widgets.get(section, {}) and value is not None:
                self.set_widget_value(section, field, value)

        # Store the quick wizard information on the detailed section tabs. Narratives are
        # generated from these fields, not from a resource-concern checklist.
        set_if("Existing condition summary", existing_summary)
        if section == "Barnyard":
            set_if("Area name / field label", condition_name or loc or "barnyard / heavy use area")
            set_if("Area size", size)
            set_if("Animal traffic / feeding details", details)
            set_if("Planned BMP", bmp)
            set_if("Planned BMP dimensions", size)
        elif section == "Clean Stormwater":
            set_if("Roof / clean water source", condition_name or "farmstead clean stormwater")
            set_if("Existing outlet / discharge location", loc)
            set_if("Current concern", existing_summary)
            set_if("Planned BMP", bmp)
            set_if("Length / size", size)
            set_if("Planned outlet route", loc)
            set_if("Additional planned notes", details)
        elif section == "Livestock Manure Storage":
            set_if("Current manure storage practice", condition_name or "manure storage and handling")
            set_if("Observed manure concern", existing_summary)
            set_if("Nearby water / sensitive area", loc)
            set_if("Planned storage BMP", bmp)
            set_if("Storage dimensions", size)
            set_if("Additional planned notes", details)
        elif section == "Pasture":
            set_if("Pasture acres / fields", size)
            set_if("Problem area", condition_name or loc)
            set_if("Resource concern", existing_summary)
            set_if("Planned BMP", bmp)
            set_if("Additional planned notes", details)
        elif section == "Animal Mortalities":
            set_if("Current mortality method", condition_name or "animal mortality management")
            set_if("Mortality site location", loc)
            set_if("Typical mortality rate", size)
            set_if("Resource concern observed", existing_summary)
            set_if("Planned / O&M action", bmp or details)
        elif section == "Fuel / Petroleum":
            set_if("Current petroleum storage", condition_name or "petroleum storage")
            set_if("Storage condition / concern", existing_summary)
            set_if("Planned / O&M action", bmp)
            set_if("Additional notes", details or loc or size)
        elif section == "Air Quality":
            set_if("Current concern", existing_summary)
            set_if("Planned / O&M action", bmp or details)
            set_if("Additional notes", loc or size)
        elif section == "Driveways":
            set_if("Current driveway condition", existing_summary)
            set_if("Observed issue", details)
            set_if("Planned / O&M action", bmp)
            set_if("Additional notes", loc or size)
        elif section == "Feed Storage":
            set_if("Current feed storage", condition_name or "feed storage")
            set_if("Storage location / method", loc)
            set_if("Observed concern", existing_summary)
            set_if("Planned / O&M action", bmp)
            set_if("Additional notes", details or size)

        # Build the section narratives from the brief existing-condition summary and selected practices.
        self.refresh_section_narratives(section, show_msg=False)

        if standards_override.strip():
            standards = standards_override
        else:
            system_hint = SECTION_SYSTEM_NAMES.get(section, "")
            standards = standards_for_bmp_choices(selected_practices, system_hint) if selected_practices.strip() else ""
        if standards.strip():
            self.set_widget_value(section, "NRCS Standards and Units Used", "Standard | Item # | Description | Units\n" + standards)
        elif concern != "Yes":
            self.set_widget_value(section, "NRCS Standards and Units Used", "Standard | Item # | Description | Units\n")

    def get_widget_value(self, section, field):
        widget = self.data_widgets.get(section, {}).get(field)
        if widget is None:
            return ""
        if hasattr(widget, "get_value"):
            return widget.get_value().strip()
        if isinstance(widget, tk.Text):
            return widget.get("1.0", "end").strip()
        return widget.get().strip()

    def set_widget_value(self, section, field, value):
        widget = self.data_widgets.get(section, {}).get(field)
        if widget is None:
            return
        if hasattr(widget, "set_value"):
            widget.set_value(value or "")
        elif isinstance(widget, tk.Text):
            widget.delete("1.0", "end")
            widget.insert("1.0", value or "")
        else:
            widget.delete(0, "end")
            widget.insert(0, value or "")

    def get_all_data(self):
        data = {}
        for section, widgets in self.data_widgets.items():
            data[section] = {}
            for key, widget in widgets.items():
                if isinstance(widget, tk.Text):
                    data[section][key] = widget.get("1.0", "end").strip()
                elif hasattr(widget, "get_value"):
                    data[section][key] = widget.get_value().strip()
                else:
                    data[section][key] = widget.get().strip()
        return data

    def set_all_data(self, data):
        for section, fields in data.items():
            if section not in self.data_widgets:
                continue
            for key, value in fields.items():
                self.set_widget_value(section, key, value)

    def clear_widget(self, widget):
        if hasattr(widget, "set_value"):
            widget.set_value("")
        elif isinstance(widget, tk.Text):
            widget.delete("1.0", "end")
        else:
            widget.delete(0, "end")

    def clear_section(self, section):
        for widget in self.data_widgets.get(section, {}).values():
            self.clear_widget(widget)

    def clear_current_section(self):
        current_index = self.notebook.index(self.notebook.select())
        section = self.tab_names[current_index] if hasattr(self, "tab_names") and current_index < len(self.tab_names) else SECTIONS[min(current_index, len(SECTIONS)-1)]
        if not messagebox.askyesno("Clear section", f"Clear all typed/generated text in {section}?"):
            return
        self.clear_section(section)

    def clear_all_text_boxes(self):
        if not messagebox.askyesno("Clear all", "Clear all typed/generated text in every section? This will not delete saved farm input files."):
            return
        for section in self.data_widgets.keys():
            self.clear_section(section)

    def clean(self, value):
        import re
        text = str(value or "").strip()
        if not text:
            return ""
        text = text.replace("&", "and")
        replacements = {
            "witner": "winter", "inclimate": "inclement", "suplimentaly": "supplementally",
            "supplementaly": "supplementally", "Mnnure": "Manure", "mnnure": "manure",
            "ino": "into", "inbetween": "between", "tnaks": "tanks",
            "mortaliy": "mortality", "approprite": "appropriate", "acces": "access",
            "landbase": "land base", "BMP's": "BMPs", "bmp's": "BMPs",
        }
        for bad, good in replacements.items():
            text = re.sub(r"\b" + re.escape(bad) + r"\b", good, text, flags=re.IGNORECASE)
        text = re.sub(r"\s+", " ", text)
        text = text.replace(" .", ".").replace(" ,", ",")
        while ".." in text:
            text = text.replace("..", ".")
        return text.strip()

    def phrase(self, value, default=""):
        value = self.clean(value)
        if not value:
            return default
        return value.rstrip(" .!?;:")

    def sentence(self, value, default=""):
        value = self.clean(value)
        return value if value else default

    def strip_placeholder(self, value):
        text = self.clean(value)
        if not text:
            return ""
        lowered = text.lower().strip(" .")
        placeholders = {"no current resource concern", "no structural bmp - continue oandm", "no structural bmp - continue o&m", "none", "n/a", "na", "[units]", "[item #]"}
        return "" if lowered in placeholders else text

    def polish_text(self, text):
        import re
        text = str(text or "")
        text = text.replace("&", "and")
        text = text.replace("County County", "County")
        text = text.replace("Basin Basin", "Basin")
        text = text.replace("Watershed Watershed", "Watershed")
        text = re.sub(r"\s+", " ", text)
        text = text.replace(" .", ".").replace(" ,", ",").replace(" ;", ";")
        while ".." in text:
            text = text.replace("..", ".")
        text = re.sub(r"\bthe the\b", "the", text, flags=re.I)
        text = re.sub(r"The farmstead is within The farmstead (?:sits across|spans) two watersheds:?\s*", "The farmstead is within ", text, flags=re.I)
        text = re.sub(r"The farmstead is within The farmstead is within", "The farmstead is within", text, flags=re.I)
        def cap(m):
            return m.group(1) + m.group(2).upper()
        text = re.sub(r"(^|[.!?]\s+)([a-z])", cap, text)
        return text.strip()

    def clean_paragraph(self, text):
        """Final proofreading pass before text is written to the report."""
        return self.proofread_text(text)

    def maybe_sentence(self, txt):
        txt = self.strip_placeholder(txt)
        if not txt:
            return ""
        txt = txt.rstrip(" ;:")
        if txt[-1] not in ".!?:":
            txt += "."
        return self.polish_text(txt)

    def join_sentences(self, parts):
        out = []
        for p in parts:
            s = self.maybe_sentence(p)
            if s:
                out.append(s)
        return self.polish_text(" ".join(out))

    def with_unit(self, value, unit):
        value = self.phrase(value)
        if not value:
            return ""
        return value if unit.lower() in value.lower() else f"{value} {unit}"

    def bullet_inventory(self, raw):
        lines = []
        for line in (raw or "").splitlines():
            line = self.clean(line).strip().strip("•-").strip()
            if line and not line.lower().startswith("example"):
                lines.append(f"• {line}")
        return "\n".join(lines)

    def clean_county(self, county):
        c = self.phrase(county, "[county]")
        if c.lower().endswith(" county"):
            c = c[:-7].strip()
        return c

    def clean_basin(self, basin):
        b = self.phrase(basin, "")
        if not b:
            return ""
        if not b.lower().endswith("basin"):
            b += " Basin"
        return b

    def clean_operation(self, operation):
        op = self.phrase(operation, "livestock operation")
        if not op.lower().endswith(("operation", "farm", "boarding")):
            op += " operation"
        return op[:1].lower() + op[1:]

    def split_items(self, value):
        text = self.clean(value).replace("\n", ";")
        if not text:
            return []
        items, seen = [], set()
        for item in text.split(";"):
            item = self.strip_placeholder(item)
            if item:
                item = item.rstrip(" .")
                key = item.lower()
                if key not in seen:
                    items.append(item)
                    seen.add(key)
        return items

    def list_phrase(self, items):
        items = [self.phrase(i) for i in items if self.phrase(i)]
        if not items:
            return ""
        if len(items) == 1:
            return items[0]
        if len(items) == 2:
            return f"{items[0]} and {items[1]}"
        return ", ".join(items[:-1]) + f", and {items[-1]}"

    def concern_text(self, value):
        items = self.split_items(value)
        if not items:
            return ""
        return "Observed resource concerns include " + self.list_phrase(items) + "."

    def practices_from_standards_text(self, raw):
        practices = []
        for line in (raw or "").splitlines():
            line = line.strip()
            if not line or line.lower().startswith("standard"):
                continue
            practice = line.split("|", 1)[0].strip()
            if practice:
                practices.append(practice)
        return practices

    def format_practices(self, text_or_standards):
        practices = self.practices_from_standards_text(text_or_standards)
        if not practices:
            practices = self.split_items(text_or_standards)
        return self.list_phrase(practices)


    def get_quick_concern(self, section):
        """Return True when the Quick Build Wizard marks a section as a concern."""
        value = self.get_widget_value("Quick Build Wizard", f"{section} - Concern").strip().lower()
        return value == "yes"

    def section_summary(self, section):
        """Use the Quick Wizard text box first, then the detailed tab summary field."""
        summary = self.get_widget_value("Quick Build Wizard", f"{section} - Existing Summary")
        if not summary:
            summary = self.f(section, "Existing condition summary")
        return self.proofread_text(summary)


    def summary_indicates_concern(self, summary):
        """Detect obvious concern wording even if the Yes/No field was left as No."""
        text = (summary or "").lower()
        if not text:
            return False
        no_patterns = ["no resource concern", "no air quality concern", "good working condition", "green and healthy", "keeps the storm water clean"]
        if any(p in text for p in no_patterns) and not any(p in text for p in ["contaminated", "comes into contact", "contact with manure", "leachate", "no secondary containment"]):
            return False
        concern_words = [
            "contaminated", "comes into contact", "contact with manure", "leachate", "manure-laden runoff",
            "no designated", "does not have", "no permanent", "temporary manure pile",
            "exposed to precipitation", "poorly drained", "secondary containment", "not covered",
            "no overfill", "mortalities are placed", "lacks", "muddy", "denuded", "bare",
            "erosion", "saturated"
        ]
        return any(word in text for word in concern_words)

    def section_practices(self, section, *fields):
        """Use the planned practice table first, then fall back to detailed-tab fields."""
        selected = self.get_widget_value("Quick Build Wizard", f"{section} - Selected Practices")
        if selected:
            return self.format_practices(selected)
        for field in fields:
            txt = self.f(section, field)
            if txt:
                return self.format_practices(txt) or self.proofread_text(txt).rstrip(".")
        return ""

    def proofread_text(self, text):
        """Deterministic proofreading pass for user-entered notes and generated paragraphs.

        This is not a language model, so it will not invent new facts. It corrects the
        common typing/automation problems that made earlier Word outputs unusable:
        duplicate punctuation, repeated phrases, poor capitalization, common typos,
        placeholders accidentally written as prose, and awkward starts such as
        "And there is" after a period.
        """
        import re
        text = str(text or "").strip()
        if not text:
            return ""
        text = text.replace("&", "and")
        replacements = {
            "witner": "winter", "inclimate": "inclement", "suplimentaly": "supplementally",
            "supplementaly": "supplementally", "suplimental": "supplemental", "Mnnure": "Manure",
            "mnnure": "manure", "ino": "into", "inbetween": "between", "tnaks": "tanks",
            "mortaliy": "mortality", "approprite": "appropriate", "acces": "access",
            "draned": "drained", "cattl": "cattle", "landbase": "land base", "drive ways": "driveways",
            "TMP": "temporary manure pile", "tmp": "temporary manure pile",
            "inclimate weather": "inclement weather", "farmstead's landbase": "farmstead's land base",
            "BMP's": "BMPs", "bmp's": "BMPs", "OandM": "O&M", "oandm": "O&M",
        }
        for bad, good in replacements.items():
            text = re.sub(r"\b" + re.escape(bad) + r"\b", good, text, flags=re.IGNORECASE)
        # Clean bullets and semicolon-style lists into normal prose when they are inside a paragraph.
        text = text.replace("\r", "\n")
        text = re.sub(r"\s*;\s*", ". ", text)
        text = re.sub(r"\s*\n+\s*", " ", text)
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"\bNo current resource concern\.?\s*", "", text, flags=re.I)
        text = re.sub(r"\bNo structural BMP - continue O&M\.?\s*", "", text, flags=re.I)
        text = text.replace("County County", "County").replace("Basin Basin", "Basin")
        text = text.replace("Watershed Watershed", "Watershed")
        text = re.sub(r"\s+([,.;:])", r"\1", text)
        text = re.sub(r"([.!?]){2,}", r"\1", text)
        text = text.replace(". .", ".").replace("..", ".")
        # Avoid sentence fragments beginning with And/But.
        text = re.sub(r"\.\s+And\s+", ". ", text)
        text = re.sub(r"\.\s+But\s+", ". However, ", text)
        text = re.sub(r"\bthe the\b", "the", text, flags=re.I)
        text = re.sub(r"\b([A-Za-z]+)\s+\1\b", r"\1", text, flags=re.I)
        # Capitalize sentence starts.
        def cap(m):
            return m.group(1) + m.group(2).upper()
        text = re.sub(r"(^|[.!?]\s+)([a-z])", cap, text.strip())
        if text and text[-1] not in ".!?":
            text += "."
        return text

    def make_existing_narrative(self, section, subject, summary, concern_yes, concern_language, no_concern_om):
        summary = self.proofread_text(summary)
        if not summary:
            summary = f"{subject} was evaluated during the farmstead assessment."
        parts = [summary]
        if concern_yes:
            parts.append(concern_language)
        else:
            parts.append("Based on the conditions observed, no significant resource concern was identified at this time.")
            parts.append(no_concern_om)
        return self.clean_paragraph(self.join_sentences(parts))

    def make_planned_narrative(self, section, subject, practices, concern_yes, loc, size, action_language, om_language):
        practices = self.proofread_text(practices).rstrip(".")
        loc = self.proofread_text(loc).rstrip(".")
        size = self.proofread_text(size).rstrip(".")
        if practices and practices.lower() not in {"no structural bmp - continue o&m", "continue o&m", "continued operation and maintenance"}:
            first = "To address these resource concerns" if concern_yes else "No significant resource concern was observed; however"
            sentence = f"{first}, it is recommended that {practices} be implemented"
            if loc:
                sentence += f" at {loc}"
            if size:
                sentence += f". Approximate planned dimensions or units include {size}"
            return self.clean_paragraph(self.join_sentences([sentence, action_language, om_language]))
        return self.clean_paragraph(self.join_sentences([
            f"No additional structural Best Management Practice is recommended for {subject.lower()} at this time because no significant resource concern was observed",
            om_language,
            "If conditions change or a resource concern develops, the area should be reevaluated and an appropriate BMP should be planned using applicable NRCS standards",
        ]))

    def refresh_all_narratives(self):
        self.refresh_front_page_narratives(show_msg=False)
        for sec in SECTIONS[1:]:
            self.refresh_section_narratives(sec, show_msg=False)
        messagebox.showinfo("Narratives refreshed", "All front-page and section narratives were generated from the short inputs. Review/edit them before creating the Word document.")

    def refresh_front_page_narratives(self, show_msg=True):
        b = {field: self.get_widget_value("Basic Info", field) for field, _ in BASIC_SHORT_FIELDS}
        b["Animal Inventory"] = self.get_widget_value("Basic Info", "Animal Inventory")

        farm = self.phrase(b.get("Farm Name"), "[Farm Name]")
        operation = self.clean_operation(b.get("Operation Type"))
        address = self.phrase(b.get("Address"), "[address]")
        town = self.phrase(b.get("Town"), "[town]")
        county = self.clean_county(b.get("County"))
        state = self.phrase(b.get("State"), "New York")
        watershed = self.phrase(clean_watershed_phrase(b.get("Watershed")), "")
        basin = self.clean_basin(b.get("Basin"))
        managed_land = self.with_unit(b.get("Managed Acres"), "acres")
        nearby = self.phrase(b.get("Nearby Location / Reference"))
        watercourse = self.clean(b.get("Main Watercourse / Drainage"))
        drainage = self.clean(b.get("Drainage / Topography Note"))
        housing = self.clean(b.get("Housing / Animal Access"))
        feed = self.clean(b.get("Feed Program"))
        manure = self.clean(b.get("Manure Handling"))
        timing = self.clean(b.get("Application Timing"))
        pasture_manure = self.clean(b.get("Pasture Manure Note"))

        intro = self.join_sentences([
            f"The following narrative addresses existing conditions and any resource concerns around the farmstead at {farm}",
            "The objective was to evaluate concentrated sources of nutrients and agricultural wastes and determine whether these materials have the potential to move off the farmstead",
            "Interviews with the farm owner, on-site evaluations, and evaluation of completed documents were used to gather information for this inventory",
            "The narrative is supported by illustrations documenting the existing conditions and planned BMPs for the farmstead",
        ])

        region = infer_ny_region(county, address)
        location = f"{farm} is a privately owned {operation} located at {address}"
        if town and town != "[town]":
            location += f", in the Town of {town}"
        if county and county != "[county]":
            location += f", {county} County"
        if state:
            location += f", {state}"
        overview_parts = [location]
        overview_parts.append(f"The facility is located within {region}")
        if watershed and basin:
            overview_parts.append(f"The farmstead is within {add_watershed_suffix(watershed)} of the {basin}")
        elif watershed:
            overview_parts.append(f"The farmstead is within {add_watershed_suffix(watershed)}")
        elif basin:
            overview_parts.append(f"The farmstead is located within the {basin}")
        if nearby:
            overview_parts.append(f"The property is {nearby}")
        if managed_land:
            overview_parts.append(f"The managed land base consists of the main farmstead and {managed_land} of farmland")
        if watercourse:
            overview_parts.append(watercourse)
        if drainage:
            overview_parts.append(drainage)
        inv = self.bullet_inventory(b.get("Animal Inventory", ""))
        overview = self.join_sentences(overview_parts)
        if inv:
            overview += "\nThe operation maintains an average year-round animal inventory of:\n" + inv

        animal_housing = self.join_sentences([
            housing,
            feed,
            "The central goal of the farm’s Comprehensive Nutrient Management Plan (CNMP) is the sustainable management of all generated nutrients to meet field needs while actively protecting local soil, water, and air resources",
        ])

        manure_management = self.join_sentences([
            manure,
            timing or "Land application occurs when weather and field conditions are suitable",
            pasture_manure,
        ])

        self.set_widget_value("Basic Info", "Introduction Narrative", intro)
        self.set_widget_value("Basic Info", "Overview Narrative", overview)
        self.set_widget_value("Basic Info", "Animal Housing and Feed Management Narrative", animal_housing)
        self.set_widget_value("Basic Info", "Manure Management and Application Narrative", manure_management)
        if show_msg:
            messagebox.showinfo("Narratives refreshed", "The front page narratives were generated from the short Basic Info inputs. Review/edit them before creating the Word document.")

    def f(self, sec, key):
        return self.get_widget_value(sec, key)

    def refresh_section_narratives(self, section, show_msg=True):
        existing, planned = self.build_narratives_for_section(section)
        self.set_widget_value(section, "Existing Conditions narrative", existing)
        self.set_widget_value(section, "Planned Actions narrative", planned)
        if show_msg:
            messagebox.showinfo("Narratives refreshed", f"{section} narratives were generated. Review/edit them before creating the Word document.")

    def use_no_resource_concern_narrative(self, section):
        existing, planned = self.build_no_resource_concern_narrative(section)
        self.set_widget_value(section, "Existing Conditions narrative", existing)
        self.set_widget_value(section, "Planned Actions narrative", planned)
        messagebox.showinfo("No resource concern narrative added", f"A shorter no-resource-concern/O&M narrative was added for {section}. Review/edit it before creating the Word document.")

    def build_no_resource_concern_narrative(self, section):
        """Short, reviewed O&M narrative used when a section has no current resource concern."""
        g = lambda k: self.strip_placeholder(self.f(section, k))
        quick_summary = g("Existing condition summary")

        if section == "Barnyard":
            name = g("Area name / field label") or "barnyard area"
            size = g("Area size")
            current = g("Current use")
            condition = g("Existing surface condition")
            details = g("Animal traffic / feeding details")
            existing = self.join_sentences([
                f"The {name}{f' ({size})' if size else ''} was evaluated during the farmstead assessment",
                quick_summary,
                f"The area is currently used for {current}" if current else "The area is used for normal livestock movement and farmstead access",
                condition,
                details,
                "Based on the evaluation, this area does not appear to be creating a significant resource concern at this time",
            ])
            planned = self.join_sentences([
                "No additional structural Best Management Practice is recommended for the barnyard at this time",
                "The farm should continue routine operation and maintenance, including maintaining stable surfaces, limiting manure accumulation, and monitoring the area during wet weather and high-use periods",
            ])
        elif section == "Clean Stormwater":
            source = g("Roof / clean water source") or "clean stormwater"
            working = g("Working gutters / downspouts note")
            existing = self.join_sentences([
                f"{source.capitalize()} was evaluated during the farmstead assessment",
                quick_summary,
                working,
                "Based on the evaluation, clean water is being managed without creating a significant resource concern at this time",
            ])
            planned = self.join_sentences([
                "No additional structural Best Management Practice is recommended for clean stormwater at this time",
                g("O&M note") or "The farm should continue to clean, inspect, and repair gutters, downspouts, outlets, and discharge areas as needed so clean water remains separated from manure and livestock use areas",
            ])
        elif section == "Livestock Manure Storage":
            practice = g("Current manure storage practice") or "current manure storage and handling practices"
            existing = self.join_sentences([
                "Livestock manure storage and handling were evaluated during the farmstead assessment",
                quick_summary,
                practice if not self.split_items(practice) or len(self.split_items(practice)) == 1 else "Current manure storage conditions were reviewed during the site visit",
                g("Application / export note"),
                g("Human waste separation note"),
                "Based on the evaluation, the current manure handling system does not appear to be creating a significant resource concern at this time",
            ])
            planned = self.join_sentences([
                "No additional structural Best Management Practice is recommended for livestock manure storage at this time",
                "The farm should continue routine operation and maintenance of manure handling areas, including monitoring for leachate, ponding, runoff, clean water contact, and any change that could allow manure or nutrients to leave the production area",
            ])
        elif section == "Pasture":
            acres = g("Pasture acres / fields")
            system = g("Current grazing system")
            details = g("Resource concern")
            existing = self.join_sentences([
                f"The pasture system{f' ({acres})' if acres else ''} was evaluated during the farmstead assessment",
                f"Current pasture management consists of {system}" if system else "Pasture areas are used for normal livestock grazing",
                details,
                "Based on the evaluation, current pasture use does not appear to be creating a significant resource concern at this time",
            ])
            planned = self.join_sentences([
                "No additional structural Best Management Practice is recommended for the pasture system at this time",
                "The farm should continue routine grazing management and maintain adequate forage cover to protect soil and water resources",
            ])
        elif section == "Animal Mortalities":
            method = g("Current mortality method") or "current animal mortality management practices"
            existing = self.join_sentences([
                "Animal mortality management was evaluated during the farmstead assessment",
                f"The farm currently manages mortalities through {method}" if method else "The farm's mortality handling practices were reviewed",
                g("Burial depth / timeframe"),
                g("Mortality site location"),
                g("Distance from watercourse"),
                g("Regulation note"),
                "Based on the evaluation, animal mortality handling does not appear to be creating a significant resource concern at this time",
            ])
            planned = self.join_sentences([
                "No additional structural Best Management Practice is recommended for animal mortality management at this time",
                g("Planned / O&M action") or "The farm should continue to follow applicable New York State requirements and maintain separation from wells, streams, ditches, wetlands, and other sensitive areas",
            ])
        elif section == "Fuel / Petroleum":
            existing = self.join_sentences([
                "Petroleum storage was evaluated during the farmstead assessment",
                g("Current petroleum storage"),
                g("Storage condition / concern"),
                g("Additional notes"),
                "Based on the evaluation, petroleum storage does not appear to be creating a significant resource concern at this time",
            ])
            planned = self.join_sentences([
                "No additional structural Best Management Practice is recommended for petroleum storage at this time",
                g("Planned / O&M action") or "Any future petroleum storage should follow applicable DEC guidance and local and state requirements",
            ])
        elif section == "Air Quality":
            existing = self.join_sentences([
                "Air quality was evaluated during the farmstead assessment",
                g("Manure type / storage"),
                g("Ammonia / methane note"),
                g("Odor condition"),
                g("Dust condition"),
                "No current air quality resource concern was observed",
            ])
            planned = self.join_sentences([
                "No structural air quality BMP is recommended at this time",
                g("Planned / O&M action") or "The farm should continue proper manure handling and application practices and remain mindful of nearby residences during windy conditions",
            ])
        elif section == "Driveways":
            existing = self.join_sentences([
                "Driveways and farm access areas were evaluated during the farmstead assessment",
                g("Current driveway condition"),
                g("Observed issue"),
                g("Additional notes"),
                "Based on the evaluation, the driveway system does not appear to be creating a significant resource concern at this time",
            ])
            planned = self.join_sentences([
                "No additional structural Best Management Practice is recommended for driveways at this time",
                g("Planned / O&M action") or "The farm should continue to maintain driveways by grading, adding gravel, repairing unstable areas, and cleaning up manure as needed",
            ])
        elif section == "Feed Storage":
            existing = self.join_sentences([
                "Feed storage was evaluated during the farmstead assessment",
                g("Current feed storage"),
                g("Storage location / method"),
                "Based on the evaluation, feed storage does not appear to be creating a significant resource concern at this time",
            ])
            planned = self.join_sentences([
                "No additional structural Best Management Practice is recommended for feed storage at this time",
                g("Planned / O&M action") or "The farm should continue routine operation and maintenance of feed storage areas and monitor for spoiled feed, leachate, runoff contact, and muddy conditions",
            ])
        else:
            existing = self.join_sentences([f"{section} was evaluated during the farmstead assessment", "No significant resource concern was observed at this time"])
            planned = self.join_sentences(["Continued operation and maintenance is recommended"])
        return existing, planned

    def build_narratives_for_section(self, s):
        """Build reviewed-style EC narratives from the Quick Wizard summary text.

        The brief existing-condition text box is the driver. The program does not
        paste checklist phrases into sentences. It proofreads the user's field
        notes, then wraps them in concise section-specific language modeled after
        completed EC evaluations: existing practice/observation, resource concern,
        planned BMP purpose, and O&M.
        """
        summary = self.section_summary(s)
        concern_yes = self.get_quick_concern(s) or self.summary_indicates_concern(summary)
        g = lambda k: self.strip_placeholder(self.f(s, k))

        if s == "Barnyard":
            name = g("Area name / field label") or self.get_widget_value("Quick Build Wizard", f"{s} - Condition Name") or "barnyard or heavy use area"
            loc = self.get_widget_value("Quick Build Wizard", f"{s} - Location") or g("Designer / agency")
            size = self.get_widget_value("Quick Build Wizard", f"{s} - Size") or g("Area size") or g("Planned BMP dimensions")
            if name and size and summary and size.lower() not in summary.lower():
                summary = f"The {name} is approximately {size}. {summary}"
            elif name and summary and name.lower() not in summary.lower():
                summary = f"The {name} was evaluated. {summary}"
            practices = self.section_practices(s, "Planned BMP")
            existing = self.make_existing_narrative(
                s,
                "Barnyard or heavy use area",
                summary,
                concern_yes,
                "This creates a resource concern because concentrated livestock traffic, exposed soil, manure accumulation, and feeding activity can allow sediment, nutrients, pathogens, and manure-laden runoff to leave the production area during rainfall or snowmelt events.",
                "Vegetative cover, surface stability, manure accumulation, and livestock traffic should continue to be monitored during wet weather and high-use periods.",
            )
            planned = self.make_planned_narrative(
                s,
                "the barnyard",
                practices,
                concern_yes,
                loc,
                size,
                "The planned work should provide a stable livestock use area, improve manure collection, reduce hoof pressure on abused areas, and limit the movement of contaminated runoff from the production area.",
                "The farm should continue to maintain stable surfaces, limit manure accumulation, clean traffic areas as needed, and monitor runoff during wet weather.",
            )
            return existing, planned

        if s == "Clean Stormwater":
            source = g("Roof / clean water source") or "clean stormwater"
            loc = self.get_widget_value("Quick Build Wizard", f"{s} - Location") or g("Planned outlet route") or g("Existing outlet / discharge location")
            size = self.get_widget_value("Quick Build Wizard", f"{s} - Size") or g("Length / size")
            if summary and source.lower() not in summary.lower():
                summary = f"Clean stormwater from {source} was evaluated. {summary}"
            practices = self.section_practices(s, "Planned BMP")
            existing = self.make_existing_narrative(
                s,
                "Clean stormwater",
                summary,
                concern_yes,
                "This creates a resource concern because clean water can become contaminated when it contacts manure, bedding, feed waste, temporary manure piles, or livestock use areas before leaving the farmstead.",
                "Existing gutters, downspouts, outlets, and discharge areas should continue to be maintained so clean water remains separated from manure, feed, and livestock use areas.",
            )
            outlet_note = g("Outlet daylight / discharge location")
            action = "The planned system should collect and convey clean water away from manure handling and livestock use areas so rainfall and snowmelt remain separated from agricultural waste."
            if outlet_note:
                action += f" The outlet should discharge at {outlet_note}, where water can leave the system without contacting manure or causing erosion."
            planned = self.make_planned_narrative(
                s,
                "clean stormwater",
                practices,
                concern_yes,
                loc,
                size,
                action,
                g("O&M note") or "The farm should continue to clean, inspect, and repair gutters, downspouts, outlets, and discharge areas as needed.",
            )
            return existing, planned

        if s == "Livestock Manure Storage":
            loc = self.get_widget_value("Quick Build Wizard", f"{s} - Location") or g("Nearby water / sensitive area")
            size = self.get_widget_value("Quick Build Wizard", f"{s} - Size") or g("Storage dimensions") or g("AWM required capacity")
            practices = self.section_practices(s, "Planned storage BMP")
            existing = self.make_existing_narrative(
                s,
                "Livestock manure storage and handling",
                summary,
                concern_yes,
                "This creates a resource concern because manure stored in temporary or uncontrolled locations can generate leachate and contaminated runoff, especially when exposed to precipitation or located near ditches, wetlands, streams, wells, poorly drained soils, or other sensitive areas.",
                "The farm should continue to monitor manure handling areas for leachate, ponding, runoff, and clean water contact.",
            )
            if concern_yes and g("Human waste separation note"):
                existing = self.clean_paragraph(existing + " " + g("Human waste separation note"))
            action = "This practice will reduce or eliminate uncontrolled temporary manure piles and allow manure to be stored in a controlled manner that is more protective of water quality."
            if g("AWM storage duration"):
                action += f" The NRCS Animal Waste Management (AWM) program should be used to size the structure for {g('AWM storage duration')} of storage."
            if g("Emptying method"):
                action += f" The storage will be emptied by {g('Emptying method')}."
            if g("Gutter / outlet / subsurface drain note"):
                action += " " + g("Gutter / outlet / subsurface drain note")
            planned = self.make_planned_narrative(
                s,
                "livestock manure storage",
                practices,
                concern_yes,
                loc,
                size,
                action,
                "The site should be evaluated by a qualified individual for soil suitability, seepage, foundation conditions, water table, and bedrock before construction.",
            )
            return existing, planned

        if s == "Pasture":
            loc = self.get_widget_value("Quick Build Wizard", f"{s} - Condition Name") or g("Problem area")
            size = self.get_widget_value("Quick Build Wizard", f"{s} - Size") or g("Pasture acres / fields")
            if loc and summary and loc.lower() not in summary.lower():
                summary = f"{loc} was evaluated as part of the pasture system. {summary}"
            if size and summary and size.lower() not in summary.lower():
                summary = f"The evaluated pasture area consists of approximately {size}. {summary}"
            practices = self.section_practices(s, "Planned BMP")
            existing = self.make_existing_narrative(
                s,
                "The pasture system",
                summary,
                concern_yes,
                "This creates a resource concern because overgrazing, poor livestock distribution, high-use areas, stream access, or lack of vegetative cover can allow soil disturbance, manure accumulation, and sediment or nutrient transport during wet weather or snowmelt events.",
                "Forage cover, livestock distribution, watering areas, laneways, and high-traffic locations should continue to be monitored throughout the grazing season.",
            )
            planned = self.make_planned_narrative(
                s,
                "the pasture system",
                practices,
                concern_yes,
                loc,
                size,
                g("Management goal") or "The planned work should improve livestock distribution, reduce grazing pressure in any one area, maintain vegetative cover, and limit sediment and nutrient transport from pasture areas.",
                "The farm should continue routine grazing management, avoid prolonged livestock concentration, and maintain adequate vegetative cover to protect soil and water resources.",
            )
            return existing, planned

        if s == "Animal Mortalities":
            loc = self.get_widget_value("Quick Build Wizard", f"{s} - Location") or g("Mortality site location")
            practices = self.section_practices(s, "Planned / O&M action")
            existing = self.make_existing_narrative(
                s,
                "Animal mortality management",
                summary,
                concern_yes,
                "This creates a resource concern because animal mortalities need a defined management method and location that prevents leachate movement, scavenger access, runoff, and impacts to wells, streams, ditches, wetlands, floodplains, property boundaries, or other sensitive areas.",
                "The mortality method, site location, handling timeframe, and separation from water resources should continue to be monitored.",
            )
            planned = self.make_planned_narrative(
                s,
                "animal mortality management",
                practices,
                concern_yes,
                loc,
                g("Typical mortality rate"),
                "The planned mortality management practice should provide a defined location and method for handling animal mortalities in a manner that limits leachate, scavenger access, runoff, and impacts to soil and water resources.",
                "Mortalities should be managed within the required timeframe and in accordance with applicable state and local requirements.",
            )
            return existing, planned

        if s == "Fuel / Petroleum":
            loc = self.get_widget_value("Quick Build Wizard", f"{s} - Location") or g("Additional notes")
            practices = self.section_practices(s, "Planned / O&M action")
            existing = self.make_existing_narrative(
                s,
                "Petroleum storage",
                summary,
                concern_yes,
                "This creates a resource concern because tanks or containers without secondary containment, overfill protection, spill response supplies, or protection from equipment traffic can allow petroleum products to reach soil, drainageways, wells, wetlands, streams, or other water resources.",
                "Tanks, containers, fueling areas, and equipment storage areas should continue to be monitored for leaks, spills, corrosion, and poor housekeeping.",
            )
            planned = self.make_planned_narrative(
                s,
                "petroleum storage",
                practices,
                concern_yes,
                loc,
                self.get_widget_value("Quick Build Wizard", f"{s} - Size"),
                "The planned practice should provide spill containment or improved storage protection to reduce the risk of petroleum products reaching soil or water resources.",
                "Operation and maintenance should include routine inspection of tanks and containers, proper housekeeping, spill response supplies, and cleanup of any staining or spilled petroleum products.",
            )
            return existing, planned

        if s == "Air Quality":
            practices = self.section_practices(s, "Planned / O&M action")
            existing = self.make_existing_narrative(
                s,
                "Air quality",
                summary,
                concern_yes,
                "This creates a resource concern when manure storage, livestock concentration areas, feed handling, driveways, or field application produce odor, dust, ammonia, methane, smoke, or nuisance conditions affecting nearby residences, roads, or other sensitive receptors.",
                "Odor, dust, manure handling, and field application conditions should continue to be monitored, especially during windy conditions and near neighboring residences or public roads.",
            )
            planned = self.make_planned_narrative(
                s,
                "air quality",
                practices,
                concern_yes,
                "",
                "",
                "The planned action should reduce odor, dust, or other nuisance conditions and support continued protection of air resources around the farmstead.",
                "The farm should continue proper manure handling and application practices, including appropriate rates, timing, and placement, and should follow all applicable DEC air regulations.",
            )
            return existing, planned

        if s == "Driveways":
            loc = self.get_widget_value("Quick Build Wizard", f"{s} - Location") or g("Additional notes")
            practices = self.section_practices(s, "Planned / O&M action")
            existing = self.make_existing_narrative(
                s,
                "Driveways and farm access areas",
                summary,
                concern_yes,
                "This creates a resource concern because rutted, muddy, unstable, or poorly drained access areas can transport manure, sediment, nutrients, or contaminated runoff from the farmstead to nearby drainageways or sensitive areas.",
                "Driveway surfaces should continue to be monitored for rutting, ponding, erosion, manure tracking, and sediment movement.",
            )
            planned = self.make_planned_narrative(
                s,
                "driveways",
                practices,
                concern_yes,
                loc,
                self.get_widget_value("Quick Build Wizard", f"{s} - Size"),
                "The planned work should improve farm access, stabilize traffic areas, and reduce the potential for sediment, manure, or nutrient-laden runoff to leave the production area.",
                "The farm should continue to maintain driveways by grading, adding gravel, repairing unstable areas, and cleaning manure from traffic areas as needed.",
            )
            return existing, planned

        if s == "Feed Storage":
            loc = self.get_widget_value("Quick Build Wizard", f"{s} - Location") or g("Storage location / method")
            practices = self.section_practices(s, "Planned / O&M action")
            existing = self.make_existing_narrative(
                s,
                "Feed storage",
                summary,
                concern_yes,
                "This creates a resource concern because silage leachate, spoiled feed, feed runoff, or clean stormwater contact with feed waste can leave the storage area or enter ditches, streams, wetlands, wells, or drainageways.",
                "Feed storage areas should continue to be monitored for spoiled feed, leachate, runoff contact, mud, wildlife issues, and traffic disturbance.",
            )
            planned = self.make_planned_narrative(
                s,
                "feed storage",
                practices,
                concern_yes,
                loc,
                self.get_widget_value("Quick Build Wizard", f"{s} - Size"),
                "The planned work should control feed runoff or leachate and prevent clean stormwater from contacting feed waste, spoiled feed, silage leachate, or muddy feeding surfaces.",
                "Spoiled feed and feed waste should be removed or managed so nutrients and organic material are not transported off-site in runoff.",
            )
            return existing, planned

        return self.clean_paragraph(summary), ""

    def save_json(self):
        path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("Farm input file", "*.json")])
        if not path:
            return
        with open(path, "w", encoding="utf-8") as f:
            json.dump(self.get_all_data(), f, indent=2)
        messagebox.showinfo("Saved", f"Farm input file saved:\n{path}")

    def load_json(self):
        path = filedialog.askopenfilename(filetypes=[("Farm input file", "*.json"), ("All files", "*.*")])
        if not path:
            return
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        self.set_all_data(data)
        messagebox.showinfo("Opened", "Farm input file loaded.")

    def set_cell_shading(self, cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

    def set_cell_text(self, cell, text, bold=False):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        if bold:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_paragraphs(self, container, text):
        text = (text or "").strip()
        if not text:
            container.add_paragraph("[No narrative entered.]")
            return
        for para in text.split("\n"):
            para = para.strip()
            if not para:
                continue
            if para.startswith("•"):
                p = container.add_paragraph(style=None)
                p.paragraph_format.left_indent = Inches(0.25)
                p.add_run(self.proofread_text(para[1:].strip()))
            else:
                container.add_paragraph(self.proofread_text(para))

    def add_section_narrative_to_cell(self, cell, section_title, narrative):
        """Place a bold section title above the narrative in a table cell."""
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(section_title)
        r.bold = True
        text = (narrative or "").strip()
        if not text:
            cell.add_paragraph("[No narrative entered.]")
            return
        for para in text.split("\n"):
            para = para.strip()
            if not para:
                continue
            if para.startswith("•"):
                bullet = cell.add_paragraph(style=None)
                bullet.paragraph_format.left_indent = Inches(0.25)
                bullet.add_run(self.proofread_text(para[1:].strip()))
            else:
                cell.add_paragraph(self.proofread_text(para))

    def parse_standards(self, raw):
        rows = []
        for line in (raw or "").splitlines():
            line = line.strip()
            if not line or line.lower().startswith("standard |") or line.lower().startswith("standard|"):
                continue
            parts = [p.strip() for p in line.split("|")]
            while len(parts) < 4:
                parts.append("")
            rows.append(parts[:4])
        return rows

    def add_standards_block(self, doc, title, raw):
        rows = self.parse_standards(raw)
        if not rows:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.bold = True
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        self.set_cell_text(hdr[0], "NRCS Standard Used:", True)
        self.set_cell_text(hdr[1], "Item # In Plan", True)
        self.set_cell_shading(hdr[0], "D9EAF7")
        self.set_cell_shading(hdr[1], "D9EAF7")
        for standard, item, desc, units in rows:
            row = table.add_row().cells
            row[0].text = ""
            p0 = row[0].paragraphs[0]
            run = p0.add_run(standard)
            run.bold = True
            if desc:
                row[0].add_paragraph(desc)
            if units:
                row[0].add_paragraph(f"Units: {units}")
            row[1].text = item
            row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        doc.add_paragraph("")

    def generate_docx(self):
        # Rebuild from the Quick Build Wizard and run the proofreading/cleanup pass
        # before the Word document is created. This prevents old messy generated
        # section text from carrying forward into the final report.
        self.refresh_front_page_narratives(show_msg=False)
        if "Quick Build Wizard" in self.data_widgets:
            for sec in QUICK_BUILD_SECTIONS:
                self.apply_quick_build_section(sec)
        else:
            for sec in SECTIONS[1:]:
                self.refresh_section_narratives(sec, show_msg=False)
        data = self.get_all_data()
        basic = data.get("Basic Info", {})

        farm_name = basic.get("Farm Name", "Farmstead") or "Farmstead"
        safe_name = "".join(c for c in farm_name if c.isalnum() or c in (" ", "_", "-")).strip().replace(" ", "_")
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"{safe_name}_Farmstead_Evaluation.docx",
            filetypes=[("Word Document", "*.docx")]
        )
        if not path:
            return

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(11)

        header = section.header.paragraphs[0]
        header.text = "WNY Crop Management Association Cooperative"
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Farmstead Evaluation-\nExisting Conditions and Planned Actions")
        run.bold = True
        run.font.size = Pt(16)
        title2 = doc.add_paragraph()
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = title2.add_run(farm_name)
        r2.bold = True
        r2.font.size = Pt(14)

        front_page = [
            ("Introduction", basic.get("Introduction Narrative", "")),
            ("Overview", basic.get("Overview Narrative", "")),
            ("Animal Housing and Feed Management", basic.get("Animal Housing and Feed Management Narrative", "")),
            ("Manure Management and Application", basic.get("Manure Management and Application Narrative", "")),
        ]
        for heading, text in front_page:
            if text.strip():
                doc.add_heading(heading, level=2)
                self.add_paragraphs(doc, text)

        eval_date = basic.get("Evaluation Date", "")
        attendees = basic.get("Attendees", "")
        if eval_date or attendees:
            sentence = "A management evaluation was conducted"
            if eval_date:
                sentence += f" on {eval_date}"
            if attendees:
                sentence += f". Attendees included {attendees}"
            sentence += "."
            doc.add_paragraph(sentence)

        doc.add_heading(f"{farm_name} Farmstead Evaluation", level=1)

        # Output each requested section in the same basic two-column format as the filled example.
        for sec in SECTIONS[1:]:
            sec_data = data.get(sec, {})
            existing = sec_data.get("Existing Conditions narrative", "")
            planned = sec_data.get("Planned Actions narrative", "")
            standards = sec_data.get("NRCS Standards and Units Used", "")

            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            cells = table.rows[0].cells
            self.set_cell_text(cells[0], "Area Evaluated- Existing Condition", True)
            self.set_cell_text(cells[1], "Area Evaluated- Planned Action", True)
            self.set_cell_shading(cells[0], "D9EAF7")
            self.set_cell_shading(cells[1], "D9EAF7")
            row = table.add_row().cells
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            self.add_section_narrative_to_cell(row[0], sec, existing)
            self.add_section_narrative_to_cell(row[1], sec, planned)
            doc.add_paragraph("")
            system_title = SECTION_SYSTEM_NAMES.get(sec, sec)
            standards_title = f"{sec} - Planned BMPs / NRCS Standards Used"
            if system_title and system_title != sec:
                standards_title += f" ({system_title})"
            self.add_standards_block(doc, standards_title, standards)

        doc.save(path)
        messagebox.showinfo("Created", f"Word document created:\n{path}")

if __name__ == "__main__":
    App().mainloop()
